import streamlit as st
# ← import 바로 다음 줄에만 이것! 다른 st.* 호출 NO
st.set_page_config(
    page_title="일상감사 접수 시스템",
    page_icon="📋",
    layout="wide",
)
from dotenv import load_dotenv  
load_dotenv()

# 이제부터 다른 import
import os
import gc  # gc 모듈 추가
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import datetime, hashlib
import requests
import json
import sqlite3
import logging
import mimetypes
import re
import ssl
import shutil
from typing import List, Dict, Optional, Tuple, Any
from docx import Document
import zipfile

# --- 페이지 상태 관리 변수 추가 (맨 위에)
if "page" not in st.session_state:
    st.session_state["page"] = "질의응답"

# 2) 여기서부터 Streamlit 호출 시작
today = datetime.datetime.now().strftime("%Y%m%d")
# 세션 쿠키 관리 추가
import uuid
if "uploader_reset_token" not in st.session_state:
    st.session_state["uploader_reset_token"] = str(uuid.uuid4())
# 앱 시작 시 새로운 세션 ID 생성
if "cookie_session_id" not in st.session_state:
    st.session_state["cookie_session_id"] = str(uuid.uuid4())
    
# submission_id 생성 시 쿠키 세션 ID 포함
if "submission_id" not in st.session_state:
    session_id = st.session_state["cookie_session_id"]
    st.session_state["submission_id"] = f"AUDIT-{today}-{session_id[:6]}"
submission_id = st.session_state["submission_id"]

# 로깅 설정
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    filename='audit_system.log'
)
logger = logging.getLogger('audit_system')

# 파일을 저장할 폴더 경로
import tempfile
base_folder = os.path.join(tempfile.gettempdir(), "uploaded_files")
if not os.path.exists(base_folder):
    os.makedirs(base_folder)

# 업로드할 날짜 정보
upload_date = datetime.datetime.now().strftime("%Y%m%d")
today_folder = os.path.join(base_folder, upload_date)
if not os.path.exists(today_folder):
    os.makedirs(today_folder)

session_folder = os.path.join(today_folder, st.session_state["submission_id"])
if not os.path.exists(session_folder):
    os.makedirs(session_folder)

# 세션 타임아웃 설정 (20분)
session_timeout = datetime.timedelta(minutes=20)

# 타임아웃 검사 및 세션 연장 로직
current_time = datetime.datetime.now()

if "last_session_time" not in st.session_state:
    # 최초 실행 시 기록
    st.session_state["last_session_time"] = current_time
    # 새 세션 시작 - 파일 업로더 상태 초기화
    for key in list(st.session_state.keys()):
        # uploader_reset_token은 남기고, 그 외 uploader_* 만 삭제
        if key.startswith("uploader_") and key != "uploader_reset_token":
            del st.session_state[key]
        # reason_ 접두사는 전부 삭제
        if key.startswith("reason_"):
            del st.session_state[key]
else:
    elapsed = current_time - st.session_state["last_session_time"]
    if elapsed > session_timeout:
        # 타임아웃 초과 시에만 세션 초기화
        keys_to_keep = ["cookie_session_id", "uploader_reset_token", "last_session_time"]
        for key in list(st.session_state.keys()):
            if key not in keys_to_keep:
                del st.session_state[key]
        # 새로운 submission_id 및 시간 갱신
        session_id = st.session_state["cookie_session_id"]
        st.session_state["submission_id"] = f"AUDIT-{today}-{session_id[:6]}"
        st.session_state["last_session_time"] = current_time
        # 임시 파일 폴더 정리
        if os.path.exists(session_folder):
            try:
                shutil.rmtree(session_folder)
                logger.info(f"세션 타임아웃으로 임시 파일 정리: {session_folder}")
            except Exception as e:
                logger.error(f"임시 파일 정리 오류: {e}")
        st.rerun()
# 정상 흐름 시 마지막 상호작용 시간 갱신
st.session_state["last_session_time"] = current_time

# ✅ GPT 감사보고서 docx 생성 함수
def generate_audit_report_with_gpt(submission_id, department, manager, phone, contract_name,
                                   contract_date, contract_amount, uploaded_files, missing_files_with_reasons) -> Optional[str]:
    try:
        # 제출 자료와 누락 자료를 읽기 쉬운 형식으로 변환
        uploaded_list = "\n".join([f"- {file}" for file in uploaded_files]) if uploaded_files else "없음"
        
        missing_list = ""
        if missing_files_with_reasons:
            missing_list = "\n".join([f"- {name}: {reason}" for name, reason in missing_files_with_reasons])
        else:
            missing_list = "없음"
        
        # 명확하고 상세한 지시사항 포함
        user_message = f"""
다음 정보를 기반으로, 상세하고 전문적인 일상감사 보고서를 작성해주세요:

## 계약 기본 정보
- 접수 ID: {submission_id}
- 접수 부서: {department}
- 담당자: {manager} (연락처: {phone})
- 계약명: {contract_name}
- 계약 체결일: {contract_date}
- 계약금액: {contract_amount}

## 제출된 자료
{uploaded_list}

## 누락된 자료 및 사유
{missing_list}


감사 전문가가 작성한 것과 같은 수준의 상세하고 전문적인 보고서를 작성해주세요.
"""
        
        # GPT 응답 가져오기
        answer, success = get_clean_answer_from_gpts(user_message)
        if not success:
            return None

        # 인용 마크 및 볼드 콜론 패턴 제거
        answer = re.sub(r'\【\d+\:\d+\†source\】', '', answer)
        answer = re.sub(r'\*\*(.*?)\:\*\*', r'\1', answer)  # **키워드:** 형태 제거
        
        document = Document()
        document.add_heading('일상감사 보고서 초안', level=0)
        
        # 보고서 내용을 적절한 형식으로 변환
        for line in answer.strip().split("\n"):
            if line.strip().startswith("# "):
                document.add_heading(line.replace("# ", "").strip(), level=1)
            elif line.strip().startswith("## "):
                document.add_heading(line.replace("## ", "").strip(), level=2)
            elif line.strip().startswith("### "):
                document.add_heading(line.replace("### ", "").strip(), level=3)
            elif line.strip().startswith("- ") or line.strip().startswith("* "):
                # 불릿 포인트 처리
                p = document.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(line.strip()[2:])
            else:
                if line.strip():  # 빈 줄이 아닌 경우만 추가
                    document.add_paragraph(line.strip())

        report_folder = os.path.join(base_folder, "draft_reports")
        os.makedirs(report_folder, exist_ok=True)
        report_path = os.path.join(report_folder, f"감사보고서초안_{submission_id}.docx")
        document.save(report_path)
        return report_path

    except Exception as e:
        logger.error(f"GPT 보고서 생성 오류: {str(e)}")
        return None

# 파일 내용 읽기 함수 추가

def extract_file_content(file_path: str) -> str:
    """
    Word와 PDF 파일의 실제 내용을 추출하여 텍스트로 반환합니다.
    """
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                return content[:3000]  # 3000자 제한
        
        elif file_ext == '.docx':
            try:
                from docx import Document
                doc = Document(file_path)
                content = []
                
                # 모든 문단 텍스트 추출
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():  # 빈 문단 제외
                        content.append(paragraph.text.strip())
                
                # 표(table) 내용도 추출
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content.append(f"[표] {cell.text.strip()}")
                
                full_content = '\n'.join(content)
                return full_content[:3000] if full_content else "[Word 파일이 비어있음]"
                
            except ImportError:
                return "[Word 파일 - python-docx 모듈 필요]"
            except Exception as e:
                return f"[Word 파일 읽기 오류: {str(e)}]"
        
        elif file_ext == '.pdf':
            try:
                import PyPDF2
                content = []
                
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    
                    # 각 페이지의 텍스트 추출
                    for page_num, page in enumerate(reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                content.append(f"[페이지 {page_num + 1}]\n{page_text.strip()}")
                        except Exception as e:
                            content.append(f"[페이지 {page_num + 1} 읽기 실패: {str(e)}]")
                
                full_content = '\n\n'.join(content)
                return full_content[:3000] if full_content else "[PDF 파일에서 텍스트 추출 실패]"
                
            except ImportError:
                return "[PDF 파일 - PyPDF2 모듈 필요]"
            except Exception as e:
                return f"[PDF 파일 읽기 오류: {str(e)}]"
        
        elif file_ext in ['.jpg', '.jpeg', '.png', '.gif']:
            return "[이미지 파일 - 텍스트 추출 불가]"
        
        elif file_ext in ['.xlsx', '.xls']:
            return "[Excel 파일 - 현재 미지원 (Word/PDF만 지원)]"
        
        else:
            # 기타 텍스트 파일 시도
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    return content[:3000]
            except UnicodeDecodeError:
                try:
                    with open(file_path, 'r', encoding='cp949') as f:
                        content = f.read()
                        return content[:3000]
                except:
                    return "[파일 내용 읽기 실패 - 인코딩 문제]"
    
    except Exception as e:
        logger.error(f"파일 내용 추출 오류: {str(e)}")
        return f"[파일 처리 중 오류 발생: {str(e)}]"

# 개선된 GPT 감사보고서 생성 함수

def generate_audit_report_with_gpt_enhanced(submission_id, department, manager, phone, contract_name,
                                           contract_date, contract_amount, uploaded_files, missing_files_with_reasons) -> Optional[str]:
    try:
        # 제출 자료의 실제 내용 추출
        uploaded_content = ""
        if uploaded_files:
            uploaded_content = "## 제출된 자료 및 내용\n\n"
            
            # DB에서 실제 파일 경로 가져오기
            conn = sqlite3.connect('audit_system.db')
            c = conn.cursor()
            
            for file_name in uploaded_files:
                c.execute("SELECT file_path FROM uploaded_files WHERE submission_id = ? AND file_name = ?", 
                         (submission_id, file_name))
                result = c.fetchone()
                
                if result and os.path.exists(result[0]):
                    file_content = extract_file_content(result[0])
                    uploaded_content += f"### 📄 {file_name}\n"
                    uploaded_content += f"```\n{file_content[:2000]}\n```\n\n"  # 내용 길이 제한
                else:
                    uploaded_content += f"### 📄 {file_name}\n[파일 내용 읽기 실패]\n\n"
            
            conn.close()
        else:
            uploaded_content = "## 제출된 자료\n없음\n\n"
        
        # 누락 자료 정리
        missing_list = ""
        if missing_files_with_reasons:
            missing_list = "## 누락된 자료 및 사유\n\n"
            missing_list += "\n".join([f"- **{name}**: {reason}" for name, reason in missing_files_with_reasons])
        else:
            missing_list = "## 누락된 자료\n없음\n\n"
        
        # 개선된 프롬프트 (실제 파일 내용 포함)
        user_message = f"""
다음 정보를 기반으로, 상세하고 전문적인 일상감사 보고서를 작성해주세요:

## 계약 기본 정보
- 접수 ID: {submission_id}
- 접수 부서: {department}
- 담당자: {manager} (연락처: {phone})
- 계약명: {contract_name}
- 계약 체결일: {contract_date}
- 계약금액: {contract_amount}

{uploaded_content}

{missing_list}

## 보고서 작성 지침
1. 제출된 파일의 실제 내용을 분석하여 구체적인 검토 의견을 제시할 것
2. 계약서, 품의서, 입찰평가표 등의 내용을 바탕으로 적정성을 평가할 것
3. 누락된 자료로 인한 제약사항을 명시할 것
4. 각 항목별로 "현황 → 검토의견 → 개선사항" 구조로 서술할 것
5. 구체적인 수치나 조건이 있다면 이를 인용하여 분석할 것
6. 전문적인 감사 관점에서 위험요소나 개선점을 도출할 것

실제 제출 자료의 내용을 기반으로 한 전문적이고 실질적인 감사보고서를 작성해주세요.
"""
        
        # GPT 응답 받기
        answer, success = get_clean_answer_from_gpts(user_message)
        if not success:
            return None

        # 인용 마크 및 볼드 콜론 패턴 제거
        answer = re.sub(r'\【\d+\:\d+\†source\】', '', answer)
        answer = re.sub(r'\*\*(.*?)\:\*\*', r'\1', answer)
        
        # Word 문서 생성
        document = Document()
        document.add_heading('일상감사 보고서 초안', level=0)
        
        # 보고서 내용을 적절한 형식으로 변환
        for line in answer.strip().split("\n"):
            if line.strip().startswith("# "):
                document.add_heading(line.replace("# ", "").strip(), level=1)
            elif line.strip().startswith("## "):
                document.add_heading(line.replace("## ", "").strip(), level=2)
            elif line.strip().startswith("### "):
                document.add_heading(line.replace("### ", "").strip(), level=3)
            elif line.strip().startswith("- ") or line.strip().startswith("* "):
                p = document.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(line.strip()[2:])
            else:
                if line.strip():
                    document.add_paragraph(line.strip())

        report_folder = os.path.join(base_folder, "draft_reports")
        os.makedirs(report_folder, exist_ok=True)
        report_path = os.path.join(report_folder, f"감사보고서초안_{submission_id}.docx")
        document.save(report_path)
        return report_path

    except Exception as e:
        logger.error(f"GPT 보고서 생성 오류: {str(e)}")
        return None

# 최적화된 GPT 감사보고서 생성 함수

def generate_audit_report_with_file_content(submission_id, department, manager, phone, contract_name,
                                           contract_date, contract_amount, uploaded_files, missing_files_with_reasons) -> Optional[str]:
    try:
        # 제출 자료의 실제 내용 추출
        uploaded_content = ""
        if uploaded_files:
            uploaded_content = "## 제출된 자료 및 실제 내용\n\n"
            
            # DB에서 실제 파일 경로 가져오기
            conn = sqlite3.connect('audit_system.db')
            c = conn.cursor()
            
            for file_name in uploaded_files:
                c.execute("SELECT file_path FROM uploaded_files WHERE submission_id = ? AND file_name LIKE ?", 
                         (submission_id, f"%{file_name.split(' - ')[0]}%"))
                result = c.fetchone()
                
                if result and os.path.exists(result[0]):
                    file_content = extract_file_content(result[0])
                    uploaded_content += f"### 📄 {file_name}\n"
                    uploaded_content += f"**파일 내용:**\n```\n{file_content}\n```\n\n"
                else:
                    uploaded_content += f"### 📄 {file_name}\n**상태:** 파일 내용 읽기 실패\n\n"
            
            conn.close()
        else:
            uploaded_content = "제출된 자료: 없음\n\n"
        
        # 누락 자료 정리
        missing_content = ""
        if missing_files_with_reasons:
            missing_content = "## 누락된 자료 및 사유\n\n"
            missing_content += "\n".join([f"- **{name}**: {reason}" for name, reason in missing_files_with_reasons])
        else:
            missing_content = "누락된 자료: 없음\n\n"
        
        # 실제 파일 내용을 포함한 프롬프트
        user_message = f"""
일상감사 보고서 초안을 작성해주세요.

## 계약 기본 정보
**접수 ID**: {submission_id}
**접수 부서**: {department}  
**담당자**: {manager} (연락처: {phone})
**계약명**: {contract_name}
**계약 체결일**: {contract_date}
**계약금액**: {contract_amount}

{uploaded_content}

{missing_content}

위의 실제 문서 내용을 분석하여 전문적인 일상감사 보고서 초안을 작성해주세요.
특히 제출된 문서의 구체적인 내용을 인용하고 분석하여 실질적인 검토 의견을 제시해주세요.
"""
        
        # GPT 응답 받기
        answer, success = get_clean_answer_from_gpts(user_message)
        if not success:
            return None

        # 보고서 파일 저장 (텍스트 파일로)
        report_folder = os.path.join(base_folder, "draft_reports")
        os.makedirs(report_folder, exist_ok=True)
        report_path = os.path.join(report_folder, f"감사보고서초안_{submission_id}.txt")
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write("일상감사 보고서 초안\n")
            f.write("=" * 50 + "\n\n")
            f.write(f"접수 ID: {submission_id}\n")
            f.write(f"접수 부서: {department}\n")  
            f.write(f"담당자: {manager} ({phone})\n")
            f.write(f"계약명: {contract_name}\n")
            f.write(f"계약 체결일: {contract_date}\n")
            f.write(f"계약금액: {contract_amount}\n\n")
            f.write("=" * 50 + "\n\n")
            f.write(answer)
        
        logger.info(f"실제 파일 내용 기반 감사보고서 생성 완료: {report_path}")
        return report_path

    except Exception as e:
        logger.error(f"GPT 보고서 생성 오류: {str(e)}")
        return None

# OpenAI API 정보 (하드코딩)
openai_api_key = st.secrets["OPENAI_API_KEY"]
openai_org_id  = st.secrets["OPENAI_ORG_ID"]

# 이메일 정보 (예시, 실제로 입력해 주세요)
from_email     = st.secrets["EMAIL_ADDRESS"]
from_password  = st.secrets["EMAIL_PASSWORD"]
to_email       = "1504282@okfngroup.com"         # 수신자 이메일 주소

# Make.com 웹훅 URL
WEBHOOK_URL = "https://hook.us2.make.com/1apecfvtsgtko5tjht4ecq3gu6qwm48v"

# 웹훅 전송 함수
def send_qa_to_webhook(session_id, question, answer, timestamp):
    """
    질의응답 데이터를 Make.com 웹훅으로 전송
    """
    try:
        payload = {
            "session_id": session_id,
            "question": question,
            "answer": answer,
            "timestamp": timestamp,
            "page": "질의응답"
        }
        
        response = requests.post(WEBHOOK_URL, json=payload, timeout=10)
        
        if response.status_code == 200:
            logger.info(f"웹훅 전송 성공: {session_id}")
            return True
        else:
            logger.error(f"웹훅 전송 실패: {response.status_code}")
            return False
            
    except Exception as e:
        logger.error(f"웹훅 전송 오류: {str(e)}")
        return False
        
# 데이터베이스 초기화
def init_db():
    try:
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        
        # 접수 내역 테이블 생성 - 확장된 필드 포함
        c.execute('''
        CREATE TABLE IF NOT EXISTS submissions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            submission_date TEXT,
            submission_id TEXT UNIQUE,
            department TEXT,
            manager TEXT,
            phone TEXT,
            contract_name TEXT,
            contract_date TEXT,
            contract_amount TEXT,
            contract_method TEXT,
            budget_item TEXT,
            status TEXT,
            email_sent INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        ''')
        
        # 기존 테이블에 새 컬럼 추가 (있으면 무시)
        try:
            c.execute("ALTER TABLE submissions ADD COLUMN contract_method TEXT")
        except sqlite3.OperationalError:
            pass  # 이미 존재하는 경우
        
        try:
            c.execute("ALTER TABLE submissions ADD COLUMN budget_item TEXT")
        except sqlite3.OperationalError:
            pass
        
        # 파일 업로드 내역 테이블 생성 (기존과 동일)
        c.execute('''
        CREATE TABLE IF NOT EXISTS uploaded_files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            submission_id TEXT,
            file_name TEXT,
            file_path TEXT,
            file_type TEXT,
            file_size INTEGER,
            uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (submission_id) REFERENCES submissions (submission_id)
        )
        ''')
        
        # 누락 파일 사유 테이블 생성 (기존과 동일)
        c.execute('''
        CREATE TABLE IF NOT EXISTS missing_file_reasons (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            submission_id TEXT,
            file_name TEXT,
            reason TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (submission_id) REFERENCES submissions (submission_id)
        )
        ''')
        
        # 인덱스 추가 (성능 최적화)
        c.execute('CREATE INDEX IF NOT EXISTS idx_submission_id ON submissions(submission_id)')
        c.execute('CREATE INDEX IF NOT EXISTS idx_submission_date ON submissions(submission_date)')
        c.execute('CREATE INDEX IF NOT EXISTS idx_uploaded_files_submission ON uploaded_files(submission_id)')
        c.execute('CREATE INDEX IF NOT EXISTS idx_missing_reasons_submission ON missing_file_reasons(submission_id)')
        
        conn.commit()
        conn.close()
        logger.info("데이터베이스 초기화 완료 (확장된 스키마)")
        return True
    except Exception as e:
        logger.error(f"데이터베이스 초기화 오류: {str(e)}")
        return False
    
# 필수 업로드 파일 목록 (누락된 파일 체크용)
required_files = [
    "계약서 파일",
    "계약 체결 관련 내부 품의서",
    "일상감사요청서",
    "입찰 평가표",
    "예산 內사용 여부",
    "업체 제안서",
    "계약 상대방 사업자등록증 또는 등기부등본",
    "소프트웨어 기술자 경력증명서 (해당할 경우)",
    "기타 관련 문서 (협약서, 과업지시서, 재무제표 등)"
]

# 파일 검증 함수 - 모든 파일 허용
def validate_file(file) -> Tuple[bool, str]:
    """
    업로드된 파일의 유효성을 검사합니다.
    모든 파일을 허용하도록 수정됨.
    
    Args:
        file: 업로드된 파일 객체
        
    Returns:
        (유효성 여부, 오류 메시지)
    """
    try:
        # 파일이 존재하는지만 확인
        if file is not None:
            return True, "파일이 유효합니다."
        return False, "파일이 없습니다."
    except Exception as e:
        logger.error(f"파일 검증 오류: {str(e)}")
        return False, f"파일 검증 중 오류가 발생했습니다: {str(e)}"

# 파일 저장 함수
def save_uploaded_file(uploaded_file, folder_path) -> Optional[str]:
    try:
        if uploaded_file is not None:
            # 파일명 보안 처리 (특수문자 제거)
            safe_filename = re.sub(r"[^\w\s.-]", "", uploaded_file.name)
            safe_filename = safe_filename.replace(" ", "_")
            
            # 세션 폴더에 저장하도록 변경
            file_path = os.path.join(session_folder, safe_filename)
            counter = 1
            while os.path.exists(file_path):
                name, ext = os.path.splitext(safe_filename)
                file_path = os.path.join(session_folder, f"{name}_{counter}{ext}")
                counter += 1
            
            # 청크 단위로 파일 저장하여 메모리 효율성 개선
            CHUNK_SIZE = 1024 * 1024  # 1MB 단위로 처리
            with open(file_path, "wb") as f:
                buffer = uploaded_file.read(CHUNK_SIZE)
                while len(buffer) > 0:
                    f.write(buffer)
                    buffer = uploaded_file.read(CHUNK_SIZE)
            
            logger.info(f"파일 저장 성공: {file_path}")
            return file_path
        return None
    except Exception as e:
        logger.error(f"파일 저장 오류: {str(e)}")
        st.error(f"파일 저장 중 오류가 발생했습니다: {str(e)}")
        return None

# 데이터베이스에 파일 정보 저장
def save_file_to_db(submission_id, file_name, file_path, file_type, file_size) -> bool:
    """
    업로드된 파일 정보를 데이터베이스에 저장합니다.
    
    Returns:
        성공 여부
    """
    try:
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        c.execute('''
        INSERT INTO uploaded_files (submission_id, file_name, file_path, file_type, file_size)
        VALUES (?, ?, ?, ?, ?)
        ''', (submission_id, file_name, file_path, file_type, file_size))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        logger.error(f"DB 파일 저장 오류: {str(e)}")
        return False

# 데이터베이스에 누락 파일 사유 저장
def save_missing_reason_to_db(submission_id, file_name, reason) -> bool:
    """
    누락된 파일 사유를 중복 없이 DB에 저장합니다.
    """
    try:
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        # 이미 같은 레코드가 있으면 삽입 안 함
        c.execute(
            "SELECT 1 FROM missing_file_reasons WHERE submission_id = ? AND file_name = ?",
            (submission_id, file_name)
        )
        if c.fetchone():
            conn.close()
            return True

        # 신규 레코드만 삽입
        c.execute('''
            INSERT INTO missing_file_reasons (submission_id, file_name, reason)
            VALUES (?, ?, ?)
        ''', (submission_id, file_name, reason))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        logger.error(f"DB 사유 저장 오류: {str(e)}")
        return False

# 데이터베이스에 접수 내역 저장 (접수 정보 포함)
def save_submission_with_info(submission_id, department, manager, phone, contract_name, contract_date, contract_amount, status="접수중", email_sent=0) -> bool:
    """
    접수 내역과 추가 정보를 데이터베이스에 저장합니다.
    
    Returns:
        성공 여부
    """
    try:
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        c.execute('''
        INSERT OR REPLACE INTO submissions
        (submission_date, submission_id, department, manager, phone, contract_name, contract_date, contract_amount, status, email_sent)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (upload_date, submission_id, department, manager, phone, contract_name, contract_date, contract_amount, status, email_sent))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        logger.error(f"DB 접수 내역 저장 오류: {str(e)}")
        return False

# 데이터베이스에서 접수 내역 업데이트
def update_submission_status(submission_id, status, email_sent=1) -> bool:
    """
    접수 내역의 상태를 업데이트합니다.
    
    Returns:
        성공 여부
    """
    try:
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        c.execute('''
        UPDATE submissions
        SET status = ?, email_sent = ?
        WHERE submission_id = ?
        ''', (status, email_sent, submission_id))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        logger.error(f"DB 접수 상태 업데이트 오류: {str(e)}")
        return False

# OpenAI API를 사용하여 질문에 답변하는 함수
def get_clean_answer_from_gpts(question: str) -> Tuple[str, bool]:
    """
    Assistant GPTs API v2 기반 GPT에게 질문을 보내고,
    최종 응답 텍스트만 추출해서 반환합니다.
    """
    try:
        import time

        assistant_id = "asst_oTip4nhZNJHinYxehJ7itwG9"

        headers = {
            "Authorization": f"Bearer {openai_api_key}",
            "OpenAI-Organization": openai_org_id,
            "Content-Type": "application/json",
            "OpenAI-Beta": "assistants=v2"
        }
        # 1. 새 스레드 생성
        thread_url = "https://api.openai.com/v1/threads"
        thread_response = requests.post(thread_url, headers=headers)
        if thread_response.status_code != 200:
            return f"[스레드 생성 실패] {thread_response.text}", False
        
        thread_id = thread_response.json()["id"]
        
        # 1. 메시지를 해당 thread에 추가
        message_url = f"https://api.openai.com/v1/threads/{thread_id}/messages"
        add_msg = {
            "role": "user",
            "content": question
        }
        msg_response = requests.post(message_url, headers=headers, json=add_msg)
        if msg_response.status_code != 200:
            return f"[메시지 추가 실패] {msg_response.text}", False

        # 2. GPT 실행 요청 (Run 생성)
        run_url = f"https://api.openai.com/v1/threads/{thread_id}/runs"
        run_response = requests.post(run_url, headers=headers, json={"assistant_id": assistant_id})
        if run_response.status_code != 200:
            return f"[실행 실패] {run_response.text}", False

        run_id = run_response.json()["id"]

        # 3. 실행 상태 확인 (폴링)
        while True:
            check = requests.get(f"{run_url}/{run_id}", headers=headers).json()
            if check["status"] == "completed":
                break
            elif check["status"] == "failed":
                return "[실행 중 실패] GPT 실행 실패", False
            time.sleep(1.5)

        # 4. 메시지 목록 조회 후 마지막 assistant 메시지의 텍스트 추출
        msgs = requests.get(message_url, headers=headers).json()["data"]
        for msg in reversed(msgs):
            if msg.get("role") == "assistant":
                for content in msg.get("content", []):
                    if content.get("type") == "text":
                        return content["text"]["value"].strip(), True

        return "[응답 없음] assistant 메시지를 찾을 수 없습니다.", False

    except Exception as e:
        return f"[예외 발생] {str(e)}", False

# OpenAI Assistant API 연동 함수
def get_assistant_response(question: str) -> str:
    """
    OpenAI Assistants API를 사용하여 질문에 대한 응답을 생성합니다.
    """
    try:
        import time
        import re  # 정규표현식 모듈 추가
        
        # 일상감사 질의응답용 Assistant ID
        assistant_id = "asst_FS7Vu9qyONYlq8O8Zab471Ek"
        
        headers = {
            "Authorization": f"Bearer {openai_api_key}",
            "OpenAI-Organization": openai_org_id,
            "Content-Type": "application/json",
            "OpenAI-Beta": "assistants=v2"
        }
        
        # 대화 맥락 유지: thread_id 세션에 저장
        if "thread_id" not in st.session_state or st.session_state.thread_id is None:
            # 새 스레드 생성
            thread_url = "https://api.openai.com/v1/threads"
            thread_response = requests.post(thread_url, headers=headers)
            if thread_response.status_code != 200:
                return f"시스템 연결에 실패했습니다. 잠시 후 다시 시도해주세요."
            thread_id = thread_response.json()["id"]
            st.session_state.thread_id = thread_id
        else:
            thread_id = st.session_state.thread_id
        
        # 메시지 추가
        message_url = f"https://api.openai.com/v1/threads/{thread_id}/messages"
        add_msg = {
            "role": "user",
            "content": question
        }
        msg_response = requests.post(message_url, headers=headers, json=add_msg)
        if msg_response.status_code != 200:
            return "메시지 전송에 실패했습니다. 다시 시도해주세요."
        
        # 스레드 실행
        run_url = f"https://api.openai.com/v1/threads/{thread_id}/runs"
        run_response = requests.post(
            run_url, 
            headers=headers, 
            json={"assistant_id": assistant_id}
        )
        if run_response.status_code != 200:
            return "처리 요청에 실패했습니다."
        
        run_id = run_response.json()["id"]
        
        # 실행 완료 확인 (폴링)
        while True:
            check = requests.get(f"{run_url}/{run_id}", headers=headers).json()
            if check["status"] == "completed":
                break
            elif check["status"] in ["failed", "cancelled", "expired"]:
                return "응답 생성에 실패했습니다. 다시 시도해주세요."
            time.sleep(1)
        
        # 메시지 목록 조회하여 응답 추출
        msgs = requests.get(message_url, headers=headers).json()["data"]
        for msg in msgs:
            if msg.get("role") == "assistant":
                for content in msg.get("content", []):
                    if content.get("type") == "text":
                        response_text = content["text"]["value"].strip()
                        # 인용 표시 제거 - 여러 형식의 인용 마크 처리
                        cleaned_response = re.sub(r'\【.*?\】', '', response_text)
                        return cleaned_response
        
        return "응답을 가져올 수 없습니다."
    
    except Exception as e:
        logger.error(f"Assistant 응답 오류: {str(e)}")
        return f"오류가 발생했습니다. 잠시 후 다시 시도해주세요."

# 이메일 발송 함수 (보안 강화)
def send_email(subject, body, to_email, attachments=None) -> Tuple[bool, str]:
    """
    이메일을 발송합니다. SSL/TLS 보안 연결을 사용합니다.
    
    Args:
        subject: 이메일 제목
        body: 이메일 본문
        to_email: 수신자 이메일
        attachments: 첨부 파일 경로 목록
        
    Returns:
        (성공 여부, 메시지)
    """
    try:
        smtp_server = "smtp.gmail.com"
        smtp_port = 465  # SSL 포트 사용
        
        msg = MIMEMultipart()
        msg["From"] = from_email
        msg["To"] = to_email
        msg["Subject"] = subject
        
        # 본문 추가
        msg.attach(MIMEText(body, "plain", "utf-8"))
        
        # 첨부 파일 추가
        if attachments:
            for file_path in attachments:
                if os.path.exists(file_path):
                    # 파일 타입 감지
                    content_type, encoding = mimetypes.guess_type(file_path)
                    if content_type is None:
                        content_type = 'application/octet-stream'
                    main_type, sub_type = content_type.split('/', 1)
                    
                    with open(file_path, "rb") as file:
                        part = MIMEApplication(file.read(), Name=os.path.basename(file_path))
                    part['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
                    msg.attach(part)
        
        # SSL 보안 연결로 SMTP 서버 연결 및 이메일 발송
        context = ssl.create_default_context()
        with smtplib.SMTP_SSL(smtp_server, smtp_port, context=context) as server:
            server.login(from_email, from_password)
            server.sendmail(from_email, to_email, msg.as_string())
        
        logger.info(f"이메일 발송 성공: {subject}")
        return True, "이메일이 성공적으로 발송되었습니다."
    except smtplib.SMTPAuthenticationError:
        error_msg = "이메일 인증 오류: 이메일 계정과 비밀번호를 확인해주세요."
        logger.error(error_msg)
        return False, error_msg
    except smtplib.SMTPException as e:
        error_msg = f"SMTP 오류: {str(e)}"
        logger.error(error_msg)
        return False, error_msg
    except Exception as e:
        error_msg = f"이메일 발송 오류: {str(e)}"
        logger.error(error_msg)
        return False, error_msg
def save_submission_with_enhanced_info(submission_id, department, manager, phone, contract_name,
                                     contract_period, contract_amount, contract_method, budget_item,
                                     status="접수중", email_sent=0) -> bool:
    """
    확장된 접수 내역과 추가 정보를 데이터베이스에 저장합니다.
    """
    try:
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()

        # 테이블에 새 컬럼이 없다면 추가 (기존 테이블 구조 확장)
        try:
            c.execute("ALTER TABLE submissions ADD COLUMN contract_method TEXT")
        except sqlite3.OperationalError:
            pass  # 컬럼이 이미 존재하는 경우

        try:
            c.execute("ALTER TABLE submissions ADD COLUMN budget_item TEXT")
        except sqlite3.OperationalError:
            pass  # 컬럼이 이미 존재하는 경우

        c.execute('''
        INSERT OR REPLACE INTO submissions
        (submission_date, submission_id, department, manager, phone, contract_name,
         contract_date, contract_amount, contract_method, budget_item, status, email_sent)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (upload_date, submission_id, department, manager, phone, contract_name,
              contract_period, contract_amount, contract_method, budget_item, status, email_sent))

        conn.commit()
        conn.close()
        return True
    except Exception as e:
        logger.error(f"DB 확장 정보 저장 오류: {str(e)}")
        return False

def generate_audit_report_with_enhanced_content(submission_id, department, manager, phone, contract_name,
                                              contract_period, contract_amount, contract_method, budget_item,
                                              uploaded_files, missing_files_with_reasons) -> Optional[str]:
    """
    1번과 2번 보고서의 장점을 결합한 확장된 접수 정보를 활용한 GPT 감사보고서 생성
    """
    try:
        # 제출 자료의 실제 내용 추출
        uploaded_content = ""
        if uploaded_files:
            uploaded_content = "## 제출된 자료 및 실제 내용\n\n"
            
            conn = sqlite3.connect('audit_system.db')
            c = conn.cursor()
            
            for file_name in uploaded_files:
                c.execute("SELECT file_path FROM uploaded_files WHERE submission_id = ? AND file_name LIKE ?", 
                         (submission_id, f"%{file_name.split(' - ')[0]}%"))
                result = c.fetchone()
                
                if result and os.path.exists(result[0]):
                    file_content = extract_file_content(result[0])
                    uploaded_content += f"### 📄 {file_name}\n"
                    uploaded_content += f"**파일 내용:**\n```\n{file_content}\n```\n\n"
                else:
                    uploaded_content += f"### 📄 {file_name}\n**상태:** 파일 내용 읽기 실패\n\n"
            
            conn.close()
        else:
            uploaded_content = "제출된 자료: 없음\n\n"
        
        # 누락 자료 정리
        missing_content = ""
        if missing_files_with_reasons:
            missing_content = "## 누락된 자료 및 사유\n\n"
            missing_content += "\n".join([f"- **{name}**: {reason}" for name, reason in missing_files_with_reasons])
        else:
            missing_content = "누락된 자료: 없음\n\n"
        
        # 더 상세하고 심층적인 분석을 위한 개선된 프롬프트
        user_message = f"""
다음 정보를 기반으로, 상세하고 전문적인 일상감사 보고서 초안을 작성해주세요.

## 📋 계약 기본 정보
**접수 ID**: {submission_id}
**접수 부서**: {department}  
**담당자**: {manager} (연락처: {phone})
**계약명**: {contract_name}
**계약 기간**: {contract_period}
**계약금액**: {contract_amount}
**계약방식**: {contract_method}
**예산과목**: {budget_item}

{uploaded_content}

{missing_content}
## 📝 상세 보고서 작성 지침

### 📏 분량 요구사항
- **전체 보고서**: 최소 2,000자 이상의 상세한 분석
- **각 검토 항목**: 최소 300-500자의 심층 분석
- **구체적 수치와 데이터**: 정확한 금액, 비율, 기간 등 명시
- **상세한 근거**: 각 판단에 대한 충분한 설명과 논리적 근거

### 🎯 구성 요구사항 - 상세 버전
**■ 사업개요** (500자 이상)
- 사업 배경과 필요성 상세 설명
- 기대효과 및 목표 구체적 기술
- 관련 법규나 정책 변화 사항 언급
- 유사 사업 사례나 벤치마킹 내용 포함

**■ 업체 선정절차** (400자 이상)
- 입찰 방식 선택 근거 상세 분석
- 참여업체 현황 및 비교 평가
- 선정 기준과 평가 방법 검토
- 절차의 투명성과 공정성 평가

**■ 검토의견: [적정/일부 부적정/부적정]** (1,200자 이상)

**1. 사업 목적 검토** (300자 이상)
- 현황 분석: 구체적 배경과 필요성
- 정당성 평가: 법적 근거와 타당성 검증
- 효과성 분석: 기대효과와 성과 지표
- 리스크 평가: 잠재적 위험요소 식별
- 개선사항: 구체적 보완 방안 제시

**2. 업체 선정 검토** (300자 이상)
- 절차 적정성: 입찰 과정의 투명성 검토
- 평가 기준: 선정 기준의 객관성과 공정성
- 법규 준수: 관련 규정 및 매뉴얼 적용 여부
- 경쟁성 확보: 충분한 경쟁 환경 조성 여부
- 개선방안: 향후 절차 개선 권고사항

**3. 예산 검토** (300자 이상)
- 예산 적정성: 배정 예산과 계약금액 비교 분석
- 집행 절차: 예산 사용 승인 과정 검토
- 비용 효율성: 유사 사업 대비 비용 분석
- 재정 영향: 조직 전체 예산에 미치는 영향
- 통제 방안: 예산 관리 및 통제 시스템

**4. 계약서 검토** (300자 이상)
- 조건 완전성: 필수 조항 포함 여부 검토
- 권리 의무: 당사자간 권리와 의무 명확성
- 리스크 관리: 위험 요소 및 대응 조항
- 변경 절차: 계약 변경 시 절차와 기준
- 분쟁 해결: 분쟁 발생 시 해결 방안

**■ 최종 의견** (400자 이상)
- 종합 평가 결과 및 근거
- 주요 개선 필요사항 우선순위별 정리
- 향후 유사 계약 시 참고사항
- 단계별 후속 조치 계획 제시

**■ 필요한 추가 자료** (구체적 목록)
- 각 자료별 제출 사유와 활용 목적 명시
- 제출 기한 및 담당자 지정 권고

### 🔍 심층 분석 요구사항
1. **정량적 분석**: 구체적 수치, 비율, 통계 데이터 활용
2. **정성적 평가**: 품질, 효과성, 적절성 등 질적 평가
3. **비교 분석**: 유사 계약, 타 기관 사례, 시장 가격 등과 비교
4. **위험 평가**: 잠재적 리스크와 대응 방안 구체적 제시
5. **개선 권고**: 단기/중기/장기별 개선 방안 상세 기술

### 🏛️ 법규 및 기준 적용
- **일상감사 매뉴얼**: 관련 조항 구체적 인용 (예: 제3.1절, 제4.2조 등)
- **계약 관련 규정**: 해당 조항과 위반 사항 명시
- **예산 관련 법규**: 예산 편성 및 집행 관련 규정 적용
- **기타 관련 법령**: 해당 업무 영역의 특별법 적용 사항

### ⚖️ 전문성 강화 원칙
- **구체적 데이터 인용**: 제출 문서의 실제 내용과 수치 활용
- **전문 용어 사용**: 감사 및 계약 관련 전문 용어 적절히 활용
- **논리적 구성**: 각 의견의 근거와 결론이 논리적으로 연결
- **실무적 관점**: 실제 업무 수행 시 고려해야 할 사항 포함

위의 지침에 따라 충분히 상세하고 전문적인 감사보고서를 작성해주세요.
간략한 요약이 아닌, 각 항목별로 심층적이고 구체적인 분석을 제공해주세요.
"""
        
        # GPT 응답 받기
        answer, success = get_clean_answer_from_gpts(user_message)
        if not success:
            return None

        # 보고서 파일 저장 (개선된 헤더 포함)
        report_folder = os.path.join(base_folder, "draft_reports")
        os.makedirs(report_folder, exist_ok=True)
        report_path = os.path.join(report_folder, f"감사보고서초안_{submission_id}.txt")
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write("일상감사 보고서 초안\n")
            f.write("=" * 80 + "\n\n")
            f.write(f"📋 접수 정보\n")
            f.write("-" * 40 + "\n")
            f.write(f"접수 ID: {submission_id}\n")
            f.write(f"접수 부서: {department}\n")  
            f.write(f"담당자: {manager} ({phone})\n")
            f.write(f"계약명: {contract_name}\n")
            f.write(f"계약 기간: {contract_period}\n")
            f.write(f"계약금액: {contract_amount}\n")
            f.write(f"계약방식: {contract_method}\n")
            f.write(f"예산과목: {budget_item}\n\n")
            f.write("=" * 80 + "\n\n")
            f.write("📝 감사 의견\n")
            f.write("-" * 40 + "\n\n")
            f.write(answer)
        
        logger.info(f"개선된 정보 기반 감사보고서 생성 완료: {report_path}")
        return report_path

    except Exception as e:
        logger.error(f"개선된 GPT 보고서 생성 오류: {str(e)}")
        return None
# 데이터베이스 초기화
init_db()

# 메뉴 정의
menu_options = ["질의응답", "파일 업로드", "접수 완료"]

# 쿼리 파라미터 대신 세션 상태 사용
menu = st.session_state["page"]

# 사이드바 메뉴
st.sidebar.title("📋 일상감사 접수 시스템")
st.sidebar.info(f"🆔 **접수 ID**: `{submission_id}`")
st.sidebar.markdown("---")

# 진행 상황 표시
progress_info = ""
if st.session_state["page"] == "질의응답":
    progress_info = "1️⃣ 질의응답 진행 중..."
elif st.session_state["page"] == "파일 업로드":
    progress_info = "2️⃣ 파일 업로드 진행 중..."
elif st.session_state["page"] == "접수 완료":
    progress_info = "3️⃣ 접수 완료 단계..."

st.sidebar.success(progress_info)
st.sidebar.markdown("---")

# 사이드바 메뉴 라디오 버튼 (아이콘 추가)
menu_options_with_icons = ["💬 질의응답", "📤 파일 업로드", "✅ 접수 완료"]
menu_mapping = {
    "💬 질의응답": "질의응답",
    "📤 파일 업로드": "파일 업로드", 
    "✅ 접수 완료": "접수 완료"
}

current_display = None
for display, actual in menu_mapping.items():
    if actual == st.session_state["page"]:
        current_display = display
        break

selected_display = st.sidebar.radio(
    "📍 메뉴 선택",
    menu_options_with_icons,
    index=menu_options_with_icons.index(current_display) if current_display else 0,
    key="menu_radio"
)

if menu_mapping[selected_display] != st.session_state["page"]:
    st.session_state["page"] = menu_mapping[selected_display]
    st.rerun()

# 사용자 가이드 (기존 위치에 개선된 내용)
with st.sidebar.expander("📖 사용 가이드", expanded=False):
    st.markdown("""
    ### 🚀 빠른 시작 가이드
    
    **1️⃣ 질의응답** *(선택사항)*
    - 🤖 AI 챗봇과 일상감사 관련 질문
    - 📋 필수 서류, 절차 등 사전 확인
    
    **2️⃣ 파일 업로드** *(필수)*
    - 📝 접수 정보 입력 (부서, 담당자 등)
    - 📁 9개 필수 서류 업로드 또는 사유 입력
    - 📊 실시간 진행률 확인
    
    **3️⃣ 접수 완료** *(필수)*
    - ✅ 최종 확인 및 이메일 발송
    - 🤖 GPT 감사보고서 초안 자동 생성
    
    ### ⚡ 주요 기능
    - 🔄 **20분 세션 타임아웃**: 보안 강화
    - 📎 **모든 파일 형식 지원**: PDF, Word, 이미지 등
    - 📧 **자동 이메일 발송**: ZIP 첨부 + 보고서
    - 🗑️ **자동 파일 정리**: 개인정보 보호
    
    ### 📞 문의처
    **OKH 감사팀**: 02-2009-6512 (신승식)
    """)

# 초기화 옵션 (기존 코드 유지)
with st.sidebar.expander("🔄 초기화 옵션", expanded=False):
    st.warning("⚠️ **주의**: 초기화 시 모든 데이터가 삭제됩니다.")
    if st.button("🔄 전체 시스템 초기화", key="btn_reset_all", use_container_width=True, type="secondary"):
        try:
            # 기존 초기화 로직 유지
            st.session_state["uploader_reset_token"] = str(uuid.uuid4())
            st.session_state["timestamp"] = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
            
            st.cache_data.clear()
            
            if os.path.exists('audit_system.db'):
                os.remove('audit_system.db')
            if os.path.exists(base_folder):
                shutil.rmtree(base_folder)
                
            keys_to_keep = ["cookie_session_id", "uploader_reset_token"]
            for key in list(st.session_state.keys()):
                if key not in keys_to_keep:
                    del st.session_state[key]
            
            session_id = st.session_state["cookie_session_id"]
            st.session_state["submission_id"] = f"AUDIT-{today}-{session_id[:6]}"
            st.session_state["last_session_time"] = datetime.datetime.now()
            
            for key in list(st.session_state.keys()):
                if key.startswith("uploader_") and key != "uploader_reset_token":
                    del st.session_state[key]
            
            st.success("✅ 시스템이 완전히 초기화되었습니다.")
            time.sleep(1)
            st.rerun()
        except Exception as e:
            st.error(f"❌ 초기화 중 오류: {e}")

# 페이지 하단 정보 개선 (기존 코드 대체)
st.sidebar.markdown("---")
st.sidebar.markdown("""
<div style="text-align: center; padding: 10px; background-color: #f0f2f6; border-radius: 5px;">
    <strong>© 2025 OKH 감사팀</strong><br>
    <small>📞 02-2009-6512 | 신승식</small><br>
    <small>🏢 일상감사 접수 시스템 v2.0</small>
</div>
""", unsafe_allow_html=True)

# 질의응답 페이지 - 첫 번째 페이지로 추가
if st.session_state["page"] == "질의응답":
    st.title("💬 일상감사 질의응답")
    
    # 헤더 정보 개선
    st.markdown("""
    <div style="background: linear-gradient(90deg, #667eea 0%, #764ba2 100%); 
                padding: 20px; border-radius: 10px; color: white; margin-bottom: 20px;">
        <h3 style="margin: 0; color: white;">🏢 OKH 일상감사 접수 시스템</h3>
        <p style="margin: 5px 0 0 0; opacity: 0.9;">질의응답 → 파일업로드 → 접수완료 순서로 진행됩니다.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # 빠른 질문 템플릿 추가
    st.markdown("### 🔥 자주 묻는 질문")
    quick_questions = [
        "📋 일상감사에 필요한 서류가 무엇인가요?",
        "💰 계약금액별 제출서류가 다른가요?",
        "📝 입찰평가표 작성 방법을 알려주세요",
        "⏰ 일상감사 처리 기간은 얼마나 걸리나요?",
        "🔄 수의계약 시 주의사항이 있나요?"
    ]
    
    cols = st.columns(3)
    for i, question in enumerate(quick_questions):
        with cols[i % 3]:
            if st.button(question, key=f"quick_q_{i}", use_container_width=True):
                # 빠른 질문을 채팅에 추가
                if "messages" not in st.session_state:
                    st.session_state.messages = []
                    st.session_state.messages.append({
                        "role": "assistant", 
                        "content": "안녕하세요! 일상감사 접수에 관해 궁금한 점을 물어봐주세요.",
                        "time": datetime.datetime.now().strftime("%H:%M")
                    })
                
                current_time = datetime.datetime.now().strftime("%H:%M")
                st.session_state.messages.append({
                    "role": "user", 
                    "content": question,
                    "time": current_time
                })
                
                # AI 응답 생성
                response = get_assistant_response(question)
                st.session_state.messages.append({
                    "role": "assistant", 
                    "content": response,
                    "time": datetime.datetime.now().strftime("%H:%M")
                })
                st.rerun()

    st.markdown("---")
    st.markdown("### 💭 AI 챗봇과 대화하기")
    
    # 세션 상태 초기화 (기존 로직 유지)
    if "messages" not in st.session_state:
        st.session_state.messages = []
        st.session_state.messages.append({
            "role": "assistant", 
            "content": "안녕하세요! 🤖 OKH 일상감사 AI 비서입니다.\n\n일상감사 접수에 관해 궁금한 점을 언제든지 물어봐주세요. 필수 서류, 절차, 작성 방법 등 모든 것을 도와드리겠습니다! 😊",
            "time": datetime.datetime.now().strftime("%H:%M")
        })
    if "thread_id" not in st.session_state:
        st.session_state.thread_id = None
    
    # 채팅 히스토리 표시 (기존 로직 유지, UI 개선)
    chat_container = st.container()
    with chat_container:
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                if message["role"] == "assistant":
                    st.markdown(f"🤖 **AI 비서** - {message['time']}")
                else:
                    st.markdown(f"👤 **나** - {message['time']}")
                st.write(message["content"])
    
    # 사용자 입력 처리 (기존 로직 유지, 플레이스홀더 추가)
    if prompt := st.chat_input("💬 궁금한 점을 입력하세요... (예: 계약서에 어떤 내용이 들어가야 하나요?)"):
        current_time = datetime.datetime.now().strftime("%H:%M")
        full_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        # 사용자 메시지 표시 및 저장
        st.session_state.messages.append({
            "role": "user", 
            "content": prompt,
            "time": current_time
        })
        with st.chat_message("user"):
            st.markdown(f"👤 **나** - {current_time}")
            st.write(prompt)

        # AI 응답 생성 중 표시
        with st.chat_message("assistant"):
            with st.spinner("🤖 AI가 답변을 생성하고 있습니다..."):
                response = get_assistant_response(prompt)
                st.markdown(f"🤖 **AI 비서** - {datetime.datetime.now().strftime('%H:%M')}")
                st.write(response)
        
        # AI 응답 저장
        st.session_state.messages.append({
            "role": "assistant", 
            "content": response,
            "time": datetime.datetime.now().strftime("%H:%M")
        })
        # ⭐ 웹훅으로 질의응답 데이터 전송 (여기에 추가)
        send_qa_to_webhook(
            session_id=submission_id,
            question=prompt,
            answer=response,
            timestamp=full_timestamp
        )
    # 채팅 통계 정보
    if len(st.session_state.messages) > 1:
        total_messages = len(st.session_state.messages) - 1  # 초기 메시지 제외
        user_messages = len([m for m in st.session_state.messages if m["role"] == "user"])
        st.caption(f"💬 대화 수: {user_messages}개 질문 | 총 {total_messages}개 메시지")
    
    st.markdown("---")
    # 다음 단계 버튼 (기존 로직 유지, UI 개선)
    next_col1, next_col2 = st.columns([3, 1])
    with next_col1:
        st.info("💡 **다음 단계**: 질문이 끝나면 파일 업로드로 진행하세요!")
    with next_col2:
        if st.button("➡️ 파일 업로드", key="next_to_upload", type="primary", use_container_width=True):
            if len(st.session_state.messages) >= 2:
                st.session_state["last_question"] = st.session_state.messages[-2]["content"]
                st.session_state["last_answer"] = st.session_state.messages[-1]["content"]
            st.session_state["page"] = "파일 업로드"
            st.rerun()

# 파일 업로드 페이지 - 완전히 수정된 버전
elif st.session_state["page"] == "파일 업로드":
    st.title("📤 일상감사 파일 업로드")

    # 📝 접수 정보 입력 - 개선된 UI
    st.markdown("### 📝 접수 정보")

    # 첫 번째 행: 기본 정보
    col1, col2 = st.columns(2)
    with col1:
        department = st.text_input(
            "🏢 접수부서",
            key="department",
            placeholder="예: IT팀, 구매팀, 총무팀"
        )
        manager = st.text_input(
            "👤 담당자",
            key="manager",
            placeholder="예: 홍길동"
        )
        phone = st.text_input(
            "📞 전화번호",
            key="phone",
            placeholder="예: 02-1234-5678 또는 010-1234-5678"
        )
    with col2:
        contract_name = st.text_input(
            "📋 계약명",
            key="contract_name",
            placeholder="예: 회계시스템 구축, ERP 도입, 홈페이지 제작"
        )

        # 계약방식 selectbox 추가
        contract_method = st.selectbox(
            "📜 계약방식",
            options=[
                "선택하세요",
                "일반경쟁입찰",
                "제한경쟁입찰",
                "지명경쟁입찰",
                "수의계약",
                "긴급계약",
                "수의시담",
                "기타"
            ],
            key="contract_method",
            help="해당 계약의 입찰 또는 계약 방식을 선택하세요"
        )

        # 예산과목 추가
        budget_item = st.text_input(
            "💰 예산과목",
            key="budget_item",
            placeholder="예: 전산개발비, 시설비, 용역비, 물품구입비"
        )

    # 두 번째 행: 계약 기간 및 금액
    st.markdown("#### 📅 계약 기간 및 금액")
    col3, col4, col5 = st.columns(3)

    with col3:
        # 계약시작일 (달력 입력)
        contract_start_date = st.date_input(
            "🗓️ 계약시작일",
            key="contract_start_date",
            help="계약 시작 예정일을 선택하세요"
        )

    with col4:
        # 계약종료일 (달력 입력)
        contract_end_date = st.date_input(
            "🗓️ 계약종료일",
            key="contract_end_date",
            help="계약 종료 예정일을 선택하세요"
        )

    with col5:
        # 계약금액
        contract_amount_str = st.text_input(
            "💵 계약금액 (원)",
            value="0",
            key="contract_amount",
            placeholder="예: 50000000 (쉼표 없이 숫자만)"
        )
        try:
            contract_amount = int(contract_amount_str.replace(',', ''))
            contract_amount_formatted = f"{contract_amount:,}"
            # 실시간 금액 표시
            if contract_amount > 0:
                st.caption(f"💡 입력금액: **{contract_amount_formatted}원**")
        except ValueError:
            contract_amount_formatted = contract_amount_str
            if contract_amount_str and contract_amount_str != "0":
                st.error("❌ 숫자만 입력해주세요 (쉼표 제외)")

    # 접수 ID 표시 및 업데이트
    if department:
        safe_dept = re.sub(r'[^\w]', '', department)[:6]
        st.session_state["submission_id"] = f"AUDIT-{upload_date}-{safe_dept}"
    sid = st.session_state["submission_id"]

    # 접수 ID를 더 눈에 띄게 표시
    st.info(f"🆔 **접수 ID**: `{sid}`")
    st.markdown("---")

    # 📋 확장된 접수 정보 DB 저장
    if all([department, manager, phone, contract_name, contract_start_date, contract_end_date, contract_amount_str, contract_method != "선택하세요"]):
        # 날짜를 문자열로 변환
        start_date_str = contract_start_date.strftime("%Y-%m-%d")
        end_date_str = contract_end_date.strftime("%Y-%m-%d")
        contract_period = f"{start_date_str} ~ {end_date_str}"

        # 확장된 정보로 DB 저장
        save_submission_with_enhanced_info(
            submission_id, department, manager, phone,
            contract_name, contract_period, contract_amount_formatted,
            contract_method, budget_item
        )

    # 📋 파일 업로드 섹션
    st.markdown("### 📋 필수 서류 업로드")
    st.markdown("💡 **안내**: 필요한 파일을 업로드하거나, 해당 파일이 없는 경우 구체적인 사유를 입력해주세요.")

    progress_container = st.container()
    progress_bar = st.progress(0)
    total_files = len(required_files)
    uploaded_count = 0

    # 파일별 아이콘 정의
    file_icons = ["📄", "📝", "🗂️", "📊", "💰", "📋", "🏢", "👨‍💻", "📁"]

    for idx, file in enumerate(required_files):
        icon = file_icons[idx] if idx < len(file_icons) else "📄"
        st.markdown(f"### {icon} {idx+1}. {file}")

        # DB에서 기존 업로드 정보 조회
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        c.execute(
            "SELECT file_name, file_path FROM uploaded_files WHERE submission_id = ? AND file_name LIKE ?",
            (submission_id, f"%{file}%")
        )
        uploaded_row = c.fetchone()
        c.execute(
            "SELECT reason FROM missing_file_reasons WHERE submission_id = ? AND file_name = ?",
            (submission_id, file)
        )
        reason_row = c.fetchone()
        conn.close()

        # 이미 업로드된 파일이 있는 경우
        if uploaded_row:
            uploaded_count += 1
            file_name, file_path = uploaded_row
            col_a, col_b = st.columns([4,1])
            with col_a:
                st.success(f"✅ **{file}** 업로드 완료: `{file_name}`")
            with col_b:
                if st.button("🗑️ 삭제", key=f"del_file_{file}"):
                    try:
                        conn = sqlite3.connect('audit_system.db')
                        c = conn.cursor()
                        c.execute(
                            "DELETE FROM uploaded_files WHERE submission_id = ? AND file_name = ?",
                            (submission_id, file_name)
                        )
                        conn.commit()
                        conn.close()
                        if os.path.exists(file_path):
                            os.remove(file_path)
                        st.success(f"🗑️ {file} 파일이 삭제되었습니다.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ 파일 삭제 오류: {e}")
            continue

        # 이미 사유가 입력된 경우
        if reason_row:
            uploaded_count += 1
            col_a, col_b = st.columns([4,1])
            with col_a:
                st.info(f"📝 **{file}** 미제출 사유: `{reason_row[0]}`")
            with col_b:
                if st.button("🗑️ 삭제", key=f"del_reason_{file}"):
                    try:
                        conn = sqlite3.connect('audit_system.db')
                        c = conn.cursor()
                        c.execute(
                            "DELETE FROM missing_file_reasons WHERE submission_id = ? AND file_name = ?",
                            (submission_id, file)
                        )
                        conn.commit()
                        conn.close()
                        st.success(f"🗑️ {file} 사유가 삭제되었습니다.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ 사유 삭제 오류: {e}")
            continue

        # 업로드 또는 사유 입력 UI
        col1, col2 = st.columns([3,1])
        with col1:
            uploaded_file = st.file_uploader(
                f"📤 **{file}** 업로드",
                key=f"uploader_{file}",
                help=f"{file}을(를) 선택하여 업로드하세요 (모든 파일 형식 지원)"
            )
        with col2:
            if uploaded_file:
                is_valid, msg = validate_file(uploaded_file)
                if is_valid:
                    path = save_uploaded_file(uploaded_file, session_folder)
                    if path:
                        save_file_to_db(
                            submission_id,
                            f"{file} - {uploaded_file.name}",
                            path,
                            os.path.splitext(uploaded_file.name)[1],
                            uploaded_file.size
                        )
                        st.success("✅ 업로드 완료")
                        uploaded_count += 1
                        st.rerun()
                else:
                    st.error(f"❌ {msg}")
            else:
                reason = st.text_input(
                    f"📝 **{file}** 미업로드 사유",
                    key=f"reason_{file}",
                    placeholder="예: 해당없음, 추후제출예정, 계약조건상 불필요",
                    help="업로드가 불가능한 구체적인 사유를 입력하세요"
                )
                if reason:
                    if save_missing_reason_to_db(submission_id, file, reason):
                        st.info("💾 사유 저장됨")
                        uploaded_count += 1
                        st.rerun()

    # 진행률 표시
    progress_percentage = uploaded_count / total_files
    progress_bar.progress(progress_percentage)

    if progress_percentage == 1.0:
        progress_container.success(f"🎉 **완료**: {uploaded_count}/{total_files} - 모든 항목이 처리되었습니다!")
    else:
        progress_container.info(f"📊 **진행상황**: {uploaded_count}/{total_files} ({progress_percentage:.1%})")

    # 다음 단계 버튼
    st.markdown("---")
    if st.button("➡️ 다음 단계: 접수 완료", type="primary", use_container_width=True):
        # 미완료 항목 체크
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        incomplete_files = []
        for req_file in required_files:
            c.execute("SELECT COUNT(*) FROM uploaded_files WHERE submission_id = ? AND file_name LIKE ?",
                      (submission_id, f"%{req_file}%"))
            file_count = c.fetchone()[0]
            c.execute("SELECT COUNT(*) FROM missing_file_reasons WHERE submission_id = ? AND file_name = ?",
                      (submission_id, req_file))
            reason_count = c.fetchone()[0]
            if file_count == 0 and reason_count == 0:
                incomplete_files.append(req_file)
        conn.close()

        if incomplete_files:
            st.warning(f"⚠️ **미완료 항목**:\n" + "\n".join([f"• {file}" for file in incomplete_files]))
        else:
            st.session_state["page"] = "접수 완료"
            st.rerun()

# 접수 완료 페이지 - 확장된 정보 처리
elif st.session_state["page"] == "접수 완료":
    st.title("✅ 일상감사 접수 완료")

    # ─── DB에서 확장된 접수 정보 불러오기 ───
    sub_id = st.session_state["submission_id"]
    conn = sqlite3.connect('audit_system.db')
    c = conn.cursor()
    
    # 확장된 컬럼 포함하여 조회
    c.execute("""
        SELECT department, manager, phone, contract_name, contract_date, contract_amount,
               contract_method, budget_item
        FROM submissions
        WHERE submission_id = ?
    """, (sub_id,))
    result = c.fetchone()
    
    if result:
        department, manager, phone, contract_name, contract_date, contract_amount, contract_method, budget_item = result
    else:
        st.error("❌ 접수 정보를 찾을 수 없습니다. 파일 업로드 페이지에서 접수 정보를 먼저 입력해주세요.")
        department, manager, phone, contract_name, contract_date, contract_amount, contract_method, budget_item = "", "", "", "", "", "", "", ""

    # 📋 접수 내용 요약 - 개선된 UI
    st.markdown("### 📋 접수 내용 요약")
    
    # 접수 정보를 카드 형태로 표시
    with st.container():
        st.markdown("""
        <div style="background-color: #f0f2f6; padding: 20px; border-radius: 10px; margin: 10px 0;">
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"""
            **🆔 접수 ID**: `{sub_id}`  
            **🏢 접수부서**: {department}  
            **👤 담당자**: {manager}  
            **📞 연락처**: {phone}  
            **📋 계약명**: {contract_name}
            """)
        with col2:
            st.markdown(f"""
            **📅 계약기간**: {contract_date}  
            **💵 계약금액**: {contract_amount}원  
            **📜 계약방식**: {contract_method}  
            **💰 예산과목**: {budget_item}  
            **📆 접수일시**: {upload_date}
            """)
        
        st.markdown("</div>", unsafe_allow_html=True)

    # 📤 업로드된 파일 목록 (기존 로직 유지, UI 개선)
    uploaded_file_list = []
    c.execute(
        "SELECT file_name, file_path FROM uploaded_files WHERE submission_id = ?",
        (sub_id,)
    )
    uploaded_db_files = c.fetchall()

    if uploaded_db_files:
        st.markdown("#### 📤 업로드된 파일")
        for idx, (file_name, file_path) in enumerate(uploaded_db_files, 1):
            st.success(f"✅ **{idx}.** {file_name}")
            uploaded_file_list.append(file_path)

    # 📝 누락된 파일 및 사유 (기존 로직 유지, UI 개선)
    c.execute(
        "SELECT file_name, reason FROM missing_file_reasons WHERE submission_id = ?",
        (sub_id,)
    )
    missing_db_files = c.fetchall()
    
    if missing_db_files:
        st.markdown("#### 📝 누락된 파일 및 사유")
        for idx, (file_name, reason) in enumerate(missing_db_files, 1):
            st.info(f"📄 **{idx}.** {file_name}: `{reason}`")

    # 완료 상태 체크 (기존 로직 유지)
    incomplete_files = []
    for req_file in required_files:
        c.execute("SELECT COUNT(*) FROM uploaded_files WHERE submission_id = ? AND file_name LIKE ?", 
                  (sub_id, f"%{req_file}%"))
        file_count = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM missing_file_reasons WHERE submission_id = ? AND file_name = ?", 
                  (sub_id, req_file))
        reason_count = c.fetchone()[0]
        if file_count == 0 and reason_count == 0:
            incomplete_files.append(req_file)
    current_missing_files = incomplete_files

    # 📧 이메일 발송 섹션 (기존 로직 유지, UI 개선)
    st.markdown("---")
    st.markdown("### 📧 이메일 발송")
    
    col1, col2 = st.columns(2)
    with col1:
        recipient_email = st.text_input(
            "📮 수신자 이메일 주소", 
            value=to_email,
            placeholder="예: audit@company.com"
        )
        email_subject = st.text_input(
            "📋 이메일 제목", 
            value=f"일상감사 접수: {submission_id}",
            placeholder="이메일 제목을 입력하세요"
        )
    with col2:
        report_recipient_email = st.text_input(
            "📬 보고서 회신 받을 이메일 주소",
            value="",
            help="감사보고서 완료 후 회신받을 이메일 주소를 입력하세요",
            placeholder="예: manager@company.com"
        )
        
    additional_message = st.text_area(
        "💬 추가 메시지", 
        value="",
        placeholder="감사팀에 전달할 추가 메시지를 입력하세요",
        height=100
    )

    # 🚀 접수 완료 버튼 (기존 로직 유지, UI 개선)
    st.markdown("---")
    if st.button('🚀 접수 완료 및 이메일 발송', type="primary", use_container_width=True):
        if current_missing_files:
            st.warning(f"⚠️ **미완료 항목**: {', '.join(current_missing_files)}. 업로드 또는 사유를 입력해 주세요.")
        else:
            with st.spinner("📤 접수 처리 중입니다..."):
                # 기존 ZIP 파일 생성 로직 유지
                zip_file_path = None
                if uploaded_file_list:
                    zip_folder = os.path.join(base_folder, "zips")
                    if not os.path.exists(zip_folder):
                        os.makedirs(zip_folder)
                    
                    zip_file_path = os.path.join(zip_folder, f"일상감사_파일_{submission_id}.zip")
                    
                    with zipfile.ZipFile(zip_file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for file_path in uploaded_file_list:
                            if os.path.exists(file_path):
                                zipf.write(file_path, os.path.basename(file_path))
                    
                    # ZIP 파일 다운로드 버튼 제공
                    with open(zip_file_path, "rb") as f:
                        zip_data = f.read()
                        st.download_button(
                            label="📥 모든 파일 다운로드 (ZIP)",
                            data=zip_data,
                            file_name=f"일상감사_파일_{submission_id}.zip",
                            mime="application/zip"
                        )
                
                # 이메일 첨부 파일 목록 준비 (기존 로직 유지)
                email_attachments = []
                if zip_file_path and os.path.exists(zip_file_path):
                    email_attachments.append(zip_file_path)
                else:
                    email_attachments.extend(uploaded_file_list)
                
                # 확장된 정보를 포함한 이메일 본문 작성
                body = f"""
📋 일상감사 접수 완료 알림

🆔 접수 ID: {submission_id}
📅 접수일자: {upload_date}
📬 보고서 회신 이메일: {report_recipient_email}

📋 접수 정보:
• 접수부서: {department}
• 담당자: {manager} ({phone})
• 계약명: {contract_name}
• 계약기간: {contract_date}
• 계약금액: {contract_amount}원
• 계약방식: {contract_method}
• 예산과목: {budget_item}

"""
                
                if additional_message:
                    body += f"\n💬 추가 메시지:\n{additional_message}\n\n"
                
                # 업로드된 파일 목록 추가 (기존 로직 유지)
                body += "📤 업로드된 파일 목록:\n"
                for file_name, _ in uploaded_db_files:
                    body += f"• {file_name}\n"
                
                # 누락된 파일 및 사유 추가 (기존 로직 유지)
                if missing_db_files:
                    body += "\n📝 누락된 파일 및 사유:\n"
                    for file_name, reason in missing_db_files:
                        body += f"• {file_name} (사유: {reason})\n"
                
                # 첨부 파일 안내 추가
                if zip_file_path:
                    body += "\n📎 업로드된 파일들이 ZIP 파일로 압축되어 첨부되어 있습니다.\n"
                
                # GPT 보고서 생성 및 첨부 (기존 로직 유지)
                report_path = generate_audit_report_with_file_content(
                    submission_id=submission_id,
                    department=st.session_state.get("department", ""),
                    manager=st.session_state.get("manager", ""),
                    phone=st.session_state.get("phone", ""),
                    contract_name=st.session_state.get("contract_name", ""),
                    contract_date=contract_date,  # 확장된 계약기간 정보
                    contract_amount=contract_amount,
                    uploaded_files=[f for f, _ in uploaded_db_files],
                    missing_files_with_reasons=[(f, r) for f, r in missing_db_files]
                )

                if report_path and os.path.exists(report_path):
                    email_attachments.append(report_path)
                    body += "🤖 GPT 기반 감사보고서 초안이 첨부되어 있습니다.\n"
                
                # 이메일 발송 (기존 로직 유지)
                success, message = send_email(email_subject, body, recipient_email, email_attachments)
                
                if success:
                    # 데이터베이스에 접수 상태 업데이트 (기존 로직 유지)
                    update_submission_status(submission_id, "접수완료", 1)
                    
                    st.success("🎉 일상감사 접수가 완료되었으며, 이메일 알림이 발송되었습니다!")
                    
                    # 접수 완료 확인서 표시 (확장된 정보 포함)
                    st.markdown("### 📄 접수 완료 확인서")
                    st.markdown(f"""
                    **🆔 접수 ID**: `{submission_id}`  
                    **📅 접수일자**: {upload_date}  
                    **🏢 접수부서**: {department}  
                    **👤 담당자**: {manager}  
                    **📋 계약명**: {contract_name}  
                    **📜 계약방식**: {contract_method}  
                    **💰 예산과목**: {budget_item}  
                    **🔄 처리상태**: 접수완료  
                    **📧 이메일 발송**: 완료 ({recipient_email})  
                    **📬 보고서 회신 이메일**: {report_recipient_email}
                    """)
                    
                    # 확장된 정보를 포함한 접수 확인서 다운로드
                    receipt_text = f"""
일상감사 접수 완료 확인서

═══════════════════════════════════════
접수 정보
═══════════════════════════════════════
접수 ID: {submission_id}
접수일자: {upload_date}
접수부서: {department}
담당자: {manager} ({phone})
계약명: {contract_name}
계약기간: {contract_date}
계약금액: {contract_amount}원
계약방식: {contract_method}
예산과목: {budget_item}

처리상태: 접수완료
이메일 발송: 완료 ({recipient_email})
보고서 회신 이메일: {report_recipient_email}

═══════════════════════════════════════
업로드된 파일 목록
═══════════════════════════════════════
"""
                    for idx, (file_name, _) in enumerate(uploaded_db_files, 1):
                        receipt_text += f"{idx}. {file_name}\n"
                    
                    if missing_db_files:
                        receipt_text += "\n═══════════════════════════════════════\n"
                        receipt_text += "누락된 파일 및 사유\n"
                        receipt_text += "═══════════════════════════════════════\n"
                        for idx, (file_name, reason) in enumerate(missing_db_files, 1):
                            receipt_text += f"{idx}. {file_name} (사유: {reason})\n"
                    
                    receipt_text += f"\n═══════════════════════════════════════\n"
                    receipt_text += f"발급일시: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                    receipt_text += f"발급기관: OKH 감사팀\n"
                    receipt_text += f"═══════════════════════════════════════"
                    
                    st.download_button(
                        label="📄 접수 확인서 다운로드",
                        data=receipt_text,
                        file_name=f"접수확인서_{submission_id}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                    
                    # 이메일 발송 후 메모리 정리 (기존 로직 유지)
                    for attachment in email_attachments:
                        if os.path.exists(attachment):
                            try:
                                # ZIP 파일은 남기고 나머지는 삭제 (선택적)
                                if not attachment.endswith('.zip'):
                                    os.remove(attachment)
                            except Exception as e:
                                logger.error(f"첨부파일 정리 오류: {str(e)}")
                    
                    # 캐시 데이터 초기화
                    st.cache_data.clear()
                    gc.collect()
                else:
                    st.error(f"❌ 이메일 발송 중 오류가 발생했습니다: {message}")

    conn.close()
