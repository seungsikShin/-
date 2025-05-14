import streamlit as st
# ← import 바로 다음 줄에만 이것! 다른 st.* 호출 NO
st.set_page_config(
    page_title="일상감사 접수 시스템",
    page_icon="📋",
    layout="wide",
)
from dotenv import load_dotenv  
load_dotenv()
with open("system_prompt.txt", "r", encoding="utf-8") as f:
    SYSTEM_PROMPT = f.read().strip()
# 이제부터 다른 import
import os
import gc  # gc 모듈 추가
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import datetime, hashlib
import requests
import json
import sqlite3
import logging
import mimetypes
import re
import ssl
import shutil
from typing import List, Dict, Optional, Tuple, Any
from docx import Document
import zipfile

# 2) 여기서부터 Streamlit 호출 시작
today = datetime.datetime.now().strftime("%Y%m%d")
# 세션 쿠키 관리 추가
import uuid
if "uploader_reset_token" not in st.session_state:
    st.session_state["uploader_reset_token"] = str(uuid.uuid4())
# 앱 시작 시 새로운 세션 ID 생성
if "cookie_session_id" not in st.session_state:
    st.session_state["cookie_session_id"] = str(uuid.uuid4())
    
# submission_id 생성 시 쿠키 세션 ID 포함
if "submission_id" not in st.session_state:
    session_id = st.session_state["cookie_session_id"]
    st.session_state["submission_id"] = f"AUDIT-{today}-{session_id[:6]}"
submission_id = st.session_state["submission_id"]

# 로깅 설정
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    filename='audit_system.log'
)
logger = logging.getLogger('audit_system')

# 파일을 저장할 폴더 경로
import tempfile
base_folder = os.path.join(tempfile.gettempdir(), "uploaded_files")
if not os.path.exists(base_folder):
    os.makedirs(base_folder)

# 업로드할 날짜 정보
upload_date = datetime.datetime.now().strftime("%Y%m%d")
today_folder = os.path.join(base_folder, upload_date)
if not os.path.exists(today_folder):
    os.makedirs(today_folder)

session_folder = os.path.join(today_folder, st.session_state["submission_id"])
if not os.path.exists(session_folder):
    os.makedirs(session_folder)

# 세션 타임아웃 설정 (20분)
session_timeout = datetime.timedelta(minutes=20)

# 타임아웃 검사 및 세션 연장 로직
current_time = datetime.datetime.now()

if "last_session_time" not in st.session_state:
    # 최초 실행 시 기록
    st.session_state["last_session_time"] = current_time
    # 새 세션 시작 - 파일 업로더 상태 초기화
    for key in list(st.session_state.keys()):
        # uploader_reset_token은 남기고, 그 외 uploader_* 만 삭제
        if key.startswith("uploader_") and key != "uploader_reset_token":
            del st.session_state[key]
        # reason_ 접두사는 전부 삭제
        if key.startswith("reason_"):
            del st.session_state[key]
else:
    elapsed = current_time - st.session_state["last_session_time"]
    if elapsed > session_timeout:
        # 타임아웃 초과 시에만 세션 초기화
        keys_to_keep = ["cookie_session_id", "uploader_reset_token", "last_session_time"]
        for key in list(st.session_state.keys()):
            if key not in keys_to_keep:
                del st.session_state[key]
        # 새로운 submission_id 및 시간 갱신
        session_id = st.session_state["cookie_session_id"]
        st.session_state["submission_id"] = f"AUDIT-{today}-{session_id[:6]}"
        st.session_state["last_session_time"] = current_time
        # 임시 파일 폴더 정리
        if os.path.exists(session_folder):
            try:
                shutil.rmtree(session_folder)
                logger.info(f"세션 타임아웃으로 임시 파일 정리: {session_folder}")
            except Exception as e:
                logger.error(f"임시 파일 정리 오류: {e}")
        st.rerun()
# 정상 흐름 시 마지막 상호작용 시간 갱신
st.session_state["last_session_time"] = current_time

# ✅ GPT 감사보고서 docx 생성 함수
def generate_audit_report_with_gpt(submission_id, department, manager, phone, contract_name,
                                   contract_date, contract_amount, uploaded_files, missing_files_with_reasons) -> Optional[str]:
    try:
        uploaded_list_str = ", ".join(uploaded_files) if uploaded_files else "없음"
        if missing_files_with_reasons:
            missing_items = "\n".join([f"- {name}: {reason}" for name, reason in missing_files_with_reasons])
        else:
            missing_items = "없음"

        user_context = f"""
        prompt = SYSTEM_PROMPT + "\n\n" + user_context
        
        answer, success = get_clean_answer_from_gpts(prompt)
        if not success:
            return None

        document = Document()
        document.add_heading('일상감사 보고서 초안', level=0)
        for line in answer.strip().split("\n"):
            if line.strip().startswith("#"):
                document.add_heading(line.replace("#", "").strip(), level=1)
            else:
                document.add_paragraph(line.strip())

        report_folder = os.path.join(base_folder, "draft_reports")
        os.makedirs(report_folder, exist_ok=True)
        report_path = os.path.join(report_folder, f"감사보고서초안_{submission_id}.docx")
        document.save(report_path)
        return report_path

    except Exception as e:
        logger.error(f"GPT 보고서 생성 오류: {str(e)}")
        return None

# OpenAI API 정보 (하드코딩)
openai_api_key = st.secrets["OPENAI_API_KEY"]
openai_org_id  = st.secrets["OPENAI_ORG_ID"]

# 이메일 정보 (예시, 실제로 입력해 주세요)
from_email     = st.secrets["EMAIL_ADDRESS"]
from_password  = st.secrets["EMAIL_PASSWORD"]
to_email       = "1504282@okfngroup.com"         # 수신자 이메일 주소


# 데이터베이스 초기화
def init_db():
    try:
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        
        # 접수 내역 테이블 생성 - 필요한 필드 추가
        c.execute('''
        CREATE TABLE IF NOT EXISTS submissions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            submission_date TEXT,
            submission_id TEXT UNIQUE,
            department TEXT,
            manager TEXT,
            phone TEXT,
            contract_name TEXT,
            contract_date TEXT,
            contract_amount TEXT,
            status TEXT,
            email_sent INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        ''')
        
        # 파일 업로드 내역 테이블 생성
        c.execute('''
        CREATE TABLE IF NOT EXISTS uploaded_files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            submission_id TEXT,
            file_name TEXT,
            file_path TEXT,
            file_type TEXT,
            file_size INTEGER,
            uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (submission_id) REFERENCES submissions (submission_id)
        )
        ''')
        
        # 누락 파일 사유 테이블 생성
        c.execute('''
        CREATE TABLE IF NOT EXISTS missing_file_reasons (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            submission_id TEXT,
            file_name TEXT,
            reason TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (submission_id) REFERENCES submissions (submission_id)
        )
        ''')
        
        conn.commit()
        conn.close()
        logger.info("데이터베이스 초기화 완료")
        return True
    except Exception as e:
        logger.error(f"데이터베이스 초기화 오류: {str(e)}")
        return False
    
# 필수 업로드 파일 목록 (누락된 파일 체크용)
required_files = [
    "계약서 파일",
    "계약 체결 관련 내부 품의서",
    "일상감사요청서",
    "입찰 평가표",
    "예산 內사용 여부",
    "업체 제안서",
    "계약 상대방 사업자등록증 또는 등기부등본",
    "소프트웨어 기술자 경력증명서 (해당할 경우)",
    "기타 관련 문서 (협약서, 과업지시서, 재무제표 등)"
]

# 파일 검증 함수 - 모든 파일 허용
def validate_file(file) -> Tuple[bool, str]:
    """
    # 업로드된 파일의 유효성을 검사합니다. (모든 파일 허용)
    
    Args:
        file: 업로드된 파일 객체
        
    Returns:
        Tuple[bool, str]: (유효성 여부, 오류 메시지)
    """
    try:
        return (True, "파일이 유효합니다.") if file else (False, "파일이 없습니다.")
    except Exception as e:
        logger.error(f"파일 검증 오류: {str(e)}")
        return False, f"검증 오류: {str(e)}"


# 파일 저장 함수
def save_uploaded_file(uploaded_file, folder_path) -> Optional[str]:
    try:
        if uploaded_file is not None:
            # 파일명 보안 처리 (특수문자 제거)
            safe_filename = re.sub(r"[^\w\s.-]", "", uploaded_file.name)
            safe_filename = safe_filename.replace(" ", "_")
            
            # 세션 폴더에 저장하도록 변경
            file_path = os.path.join(session_folder, safe_filename)
            counter = 1
            while os.path.exists(file_path):
                name, ext = os.path.splitext(safe_filename)
                file_path = os.path.join(session_folder, f"{name}_{counter}{ext}")
                counter += 1
            
            # 청크 단위로 파일 저장하여 메모리 효율성 개선
            CHUNK_SIZE = 1024 * 1024  # 1MB 단위로 처리
            with open(file_path, "wb") as f:
                buffer = uploaded_file.read(CHUNK_SIZE)
                while len(buffer) > 0:
                    f.write(buffer)
                    buffer = uploaded_file.read(CHUNK_SIZE)
            
            logger.info(f"파일 저장 성공: {file_path}")
            return file_path
        return None
    except Exception as e:
        logger.error(f"파일 저장 오류: {str(e)}")
        st.error(f"파일 저장 중 오류가 발생했습니다: {str(e)}")
        return None

# 데이터베이스에 파일 정보 저장
def save_file_to_db(submission_id, file_name, file_path, file_type, file_size) -> bool:
    """
    업로드된 파일 정보를 데이터베이스에 저장합니다.
    
    Returns:
        성공 여부
    """
    try:
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        c.execute('''
        INSERT INTO uploaded_files (submission_id, file_name, file_path, file_type, file_size)
        VALUES (?, ?, ?, ?, ?)
        ''', (submission_id, file_name, file_path, file_type, file_size))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        logger.error(f"DB 파일 저장 오류: {str(e)}")
        return False

# 데이터베이스에 누락 파일 사유 저장
def save_missing_reason_to_db(submission_id, file_name, reason) -> bool:
    """
    누락된 파일 사유를 중복 없이 DB에 저장합니다.
    """
    try:
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        # 이미 같은 레코드가 있으면 삽입 안 함
        c.execute(
            "SELECT 1 FROM missing_file_reasons WHERE submission_id = ? AND file_name = ?",
            (submission_id, file_name)
        )
        if c.fetchone():
            conn.close()
            return True

        # 신규 레코드만 삽입
        c.execute('''
            INSERT INTO missing_file_reasons (submission_id, file_name, reason)
            VALUES (?, ?, ?)
        ''', (submission_id, file_name, reason))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        logger.error(f"DB 사유 저장 오류: {str(e)}")
        return False

# 데이터베이스에 접수 내역 저장 (접수 정보 포함)
def save_submission_with_info(submission_id, department, manager, phone, contract_name, contract_date, contract_amount, status="접수중", email_sent=0) -> bool:
    """
    접수 내역과 추가 정보를 데이터베이스에 저장합니다.
    
    Returns:
        성공 여부
    """
    try:
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        c.execute('''
        INSERT OR REPLACE INTO submissions
        (submission_date, submission_id, department, manager, phone, contract_name, contract_date, contract_amount, status, email_sent)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (upload_date, submission_id, department, manager, phone, contract_name, contract_date, contract_amount, status, email_sent))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        logger.error(f"DB 접수 내역 저장 오류: {str(e)}")
        return False

# 데이터베이스에서 접수 내역 업데이트
def update_submission_status(submission_id, status, email_sent=1) -> bool:
    """
    접수 내역의 상태를 업데이트합니다.
    
    Returns:
        성공 여부
    """
    try:
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        c.execute('''
        UPDATE submissions
        SET status = ?, email_sent = ?
        WHERE submission_id = ?
        ''', (status, email_sent, submission_id))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        logger.error(f"DB 접수 상태 업데이트 오류: {str(e)}")
        return False

# OpenAI API를 사용하여 질문에 답변하는 함수
def get_clean_answer_from_gpts(question: str) -> Tuple[str, bool]:
    """
    Assistant GPTs API v2 기반 GPT에게 질문을 보내고,
    최종 응답 텍스트만 추출해서 반환합니다.
    """
    try:
        import time

        assistant_id = "asst_oTip4nhZNJHinYxehJ7itwG9"

        headers = {
            "Authorization": f"Bearer {openai_api_key}",
            "OpenAI-Organization": openai_org_id,
            "Content-Type": "application/json",
            "OpenAI-Beta": "assistants=v2"
        }
        # 1. 새 스레드 생성
        thread_url = "https://api.openai.com/v1/threads"
        thread_response = requests.post(thread_url, headers=headers)
        if thread_response.status_code != 200:
            return f"[스레드 생성 실패] {thread_response.text}", False
        
        thread_id = thread_response.json()["id"]
        
        # 1. 메시지를 해당 thread에 추가
        message_url = f"https://api.openai.com/v1/threads/{thread_id}/messages"
        add_msg = {
            "role": "user",
            "content": question
        }
        msg_response = requests.post(message_url, headers=headers, json=add_msg)
        if msg_response.status_code != 200:
            return f"[메시지 추가 실패] {msg_response.text}", False

        # 2. GPT 실행 요청 (Run 생성)
        run_url = f"https://api.openai.com/v1/threads/{thread_id}/runs"
        run_response = requests.post(run_url, headers=headers, json={"assistant_id": assistant_id})
        if run_response.status_code != 200:
            return f"[실행 실패] {run_response.text}", False

        run_id = run_response.json()["id"]

        # 3. 실행 상태 확인 (폴링)
        while True:
            check = requests.get(f"{run_url}/{run_id}", headers=headers).json()
            if check["status"] == "completed":
                break
            elif check["status"] == "failed":
                return "[실행 중 실패] GPT 실행 실패", False
            time.sleep(1.5)

        # 4. 메시지 목록 조회 후 마지막 assistant 메시지의 텍스트 추출
        msgs = requests.get(message_url, headers=headers).json()["data"]
        for msg in reversed(msgs):
            if msg.get("role") == "assistant":
                for content in msg.get("content", []):
                    if content.get("type") == "text":
                        return content["text"]["value"].strip(), True

        return "[응답 없음] assistant 메시지를 찾을 수 없습니다.", False

    except Exception as e:
        return f"[예외 발생] {str(e)}", False
        
# 이메일 발송 함수 (보안 강화)
def send_email(subject, body, to_email, attachments=None) -> Tuple[bool, str]:
    """
    이메일을 발송합니다. SSL/TLS 보안 연결을 사용합니다.
    
    Args:
        subject: 이메일 제목
        body: 이메일 본문
        to_email: 수신자 이메일
        attachments: 첨부 파일 경로 목록
        
    Returns:
        (성공 여부, 메시지)
    """
    try:
        smtp_server = "smtp.gmail.com"
        smtp_port = 465  # SSL 포트 사용
        
        msg = MIMEMultipart()
        msg["From"] = from_email
        msg["To"] = to_email
        msg["Subject"] = subject
        
        # 본문 추가
        msg.attach(MIMEText(body, "plain", "utf-8"))
        
        # 첨부 파일 추가
        if attachments:
            for file_path in attachments:
                if os.path.exists(file_path):
                    # 파일 타입 감지
                    content_type, encoding = mimetypes.guess_type(file_path)
                    if content_type is None:
                        content_type = 'application/octet-stream'
                    main_type, sub_type = content_type.split('/', 1)
                    
                    with open(file_path, "rb") as file:
                        part = MIMEApplication(file.read(), Name=os.path.basename(file_path))
                    part['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
                    msg.attach(part)
        
        # SSL 보안 연결로 SMTP 서버 연결 및 이메일 발송
        context = ssl.create_default_context()
        with smtplib.SMTP_SSL(smtp_server, smtp_port, context=context) as server:
            server.login(from_email, from_password)
            server.sendmail(from_email, to_email, msg.as_string())
        
        logger.info(f"이메일 발송 성공: {subject}")
        return True, "이메일이 성공적으로 발송되었습니다."
    except smtplib.SMTPAuthenticationError:
        error_msg = "이메일 인증 오류: 이메일 계정과 비밀번호를 확인해주세요."
        logger.error(error_msg)
        return False, error_msg
    except smtplib.SMTPException as e:
        error_msg = f"SMTP 오류: {str(e)}"
        logger.error(error_msg)
        return False, error_msg
    except Exception as e:
        error_msg = f"이메일 발송 오류: {str(e)}"
        logger.error(error_msg)
        return False, error_msg

# 데이터베이스 초기화
init_db()

# 메뉴 정의
menu_options = ["파일 업로드", "접수 완료"]

# 쿼리 파라미터에서 메뉴 초기값 가져오기
default_menu = st.query_params.get("menu", "파일 업로드")
if isinstance(default_menu, list):
    default_menu = default_menu[0]
if default_menu not in menu_options:
    default_menu = "파일 업로드"
  
# 사이드바 메뉴
st.sidebar.title("📋 일상감사 접수 시스템")
st.sidebar.info(f"접수 ID: {submission_id}")
st.sidebar.markdown("---")

with st.sidebar.expander("초기화 옵션", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        if st.button("새 접수 시작", key="btn_new_submission"):
            st.session_state["uploader_reset_token"] = str(uuid.uuid4())
            # 타임스탬프 갱신
            st.session_state["timestamp"] = datetime.datetime.now().strftime("%Y%m%d%H%M%S")     
            # 세션 상태 초기화 (쿠키 ID와 타임스탬프 제외)
            keys_to_keep = ["cookie_session_id", "uploader_reset_token"]
            for key in list(st.session_state.keys()):
                if key not in keys_to_keep:
                    del st.session_state[key]

            # 새로운 submission_id 생성
            session_id = st.session_state["cookie_session_id"]
            st.session_state["submission_id"] = f"AUDIT-{today}-{session_id[:6]}"
            st.session_state["last_session_time"] = datetime.datetime.now()
            st.success("새 접수가 시작되었습니다.")
            st.rerun()
    with col2:
        if st.button("DB 및 파일 완전 초기화", key="btn_full_reset"):
            st.session_state["uploader_reset_token"] = str(uuid.uuid4())
            try:
                if os.path.exists('audit_system.db'):
                    os.remove('audit_system.db')
                if os.path.exists(base_folder):
                    shutil.rmtree(base_folder)
                st.success("DB 및 파일 시스템이 완전히 초기화되었습니다. 새로고침 해주세요!")
                st.rerun()
            except Exception as e:
                st.error(f"오류: {e}")

    # 새로운 버튼 추가
    if st.button("파일 업로더 캐시 초기화", key="btn_clear_uploader"):
        st.cache_data.clear()
        # 파일 업로더 관련 세션 상태 변수 초기화
        for key in list(st.session_state.keys()):
            if key.startswith("uploader_") and key != "uploader_reset_token":
                del st.session_state[key]
        st.success("파일 업로더 캐시가 초기화되었습니다.")
        st.rerun()


# 메뉴 선택 라디오 버튼 (쿼리 파라미터 기반 index 설정)
menu = st.sidebar.radio(
    "메뉴 선택",
    menu_options,
    index=menu_options.index(default_menu),
    key="menu"
)


# 파일 업로드 페이지 - menu 변수가 정의된 후에 사용
if menu == "파일 업로드":
    st.title("📤 일상감사 파일 업로드")

    # 접수 정보 입력 섹션 추가
    st.markdown("### 접수 정보")
    
    # 두 개의 열로 나누어 정보 입력 필드 배치
    col1, col2 = st.columns(2)
    
    with col1:
        department = st.text_input("접수부서", key="department")
        manager = st.text_input("담당자", key="manager")
        phone = st.text_input("전화번호", key="phone")
    
    with col2:
        contract_name = st.text_input("계약명", key="contract_name")
        contract_date = st.text_input("계약 체결일(예상)", key="contract_date")
        
        # 계약금액 입력 (텍스트 입력으로 변경)
        contract_amount_str = st.text_input("계약금액", value="0", key="contract_amount")
        
        # 쉼표 제거 후 숫자로 변환 시도
        try:
            contract_amount = int(contract_amount_str.replace(',', ''))
            # 다시 형식화하여 저장
            contract_amount_formatted = f"{contract_amount:,}"
        except ValueError:
            if contract_amount_str:
                st.error("계약금액은 숫자만 입력해주세요.")
            contract_amount_formatted = contract_amount_str
    
    # 접수 ID 생성 - 부서명 포함
    if department:
        # 부서명의 첫 글자만 추출하여 ID에 포함
        safe_dept = re.sub(r'[^\w]', '', department)[:6]
        st.session_state["submission_id"] = f"AUDIT-{upload_date}-{safe_dept}"
    
    # 접수 ID 표시
    sid = st.session_state.get("submission_id", submission_id)
    st.info(f"접수 ID: {sid}")
    st.markdown("---")
    
    # 접수 정보 저장
    if all([department, manager, phone, contract_name, contract_date, contract_amount_str]):
    # 데이터 저장
        save_submission_with_info(
            submission_id,
            department,
            manager,
            phone,
            contract_name,
            contract_date,
            contract_amount_formatted
        )
      
    # 필요한 파일을 업로드하거나 사유 입력 안내
    st.markdown("필요한 파일을 업로드하거나, 해당 파일이 없는 경우 사유를 입력해주세요.")
    
    # 진행 상황 표시
    progress_container = st.container()
    progress_bar = st.progress(0)
    total_files = len(required_files)
    uploaded_count = 0
    
    # 각 파일에 대한 업로드 칸을 생성하고 체크 표시 및 사유 입력 받기
    for idx, file in enumerate(required_files):
        st.markdown(f"### {idx+1}. {file}")
        col1, col2 = st.columns([3, 1])
        
        # 파일 유형 별 DB에 업로드됐는지 확인
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        c.execute("SELECT file_name FROM uploaded_files WHERE submission_id = ? AND file_name LIKE ?", 
                (submission_id, f"%{file}%"))
        is_file_uploaded = bool(c.fetchone())
        
        # 사유 입력됐는지 확인
        c.execute("SELECT reason FROM missing_file_reasons WHERE submission_id = ? AND file_name = ?", 
                (submission_id, file))
        reason_record = c.fetchone()
        conn.close()
        
        # 이미 업로드된 파일이면 메시지만 표시
        if is_file_uploaded:
            st.success(f"✅ {file} 업로드 완료됨")
            uploaded_count += 1
            continue
        
        # 이미 사유가 있는 경우 표시
        if reason_record:
            st.info(f"📝 {file}: {reason_record[0]}")
            uploaded_count += 1
            continue
        
        with col1:
            # 사용자별 고유 키 생성
            user_key = st.session_state["cookie_session_id"]
            if "timestamp" not in st.session_state:
                st.session_state["timestamp"] = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
            timestamp = st.session_state["timestamp"]
            
            # 파일 업로더에 사용자별 고유 키 사용
            uploaded_file = st.file_uploader(
                f"📄 {file} 업로드", 
                type=None,
                key=f"uploader_{st.session_state['uploader_reset_token']}_{file}"
            )

        with col2:
            if uploaded_file:
                # 파일 검증
                is_valid, message = validate_file(uploaded_file)
        
                if is_valid:
                    # 파일 저장
                    file_path = save_uploaded_file(uploaded_file, session_folder)

                    if file_path:
                        # 파일 정보와 필수 파일 유형 정보도 함께 저장
                        file_type = os.path.splitext(uploaded_file.name)[1]
                        save_file_to_db(
                            submission_id, 
                            f"{file} - {uploaded_file.name}", # 파일 유형을 파일명에 포함
                            file_path, 
                            file_type, 
                            uploaded_file.size
                        )
                        st.success(f"✅ 업로드 완료")
                        uploaded_count += 1
                        
                        # 메모리 해제를 위한 코드 추가
                        del uploaded_file
                        gc.collect()
                        
                        # 페이지 다시 로드하여 UI 갱신
                        st.rerun()
                else:
                    st.error(message)
            else:
                reason = st.text_input(
                    f"{file} 업로드하지 않은 이유", 
                    key=f"reason_{user_key}_{timestamp}_{file}",
                    help="파일을 업로드하지 않는 경우 반드시 사유를 입력해주세요."
                )
                
                if reason:
                    if save_missing_reason_to_db(submission_id, file, reason):
                        st.info("사유가 저장되었습니다.")
                        uploaded_count += 1
                        # 사유 저장 후 페이지 리로드
                        st.rerun()

    st.markdown("---")

    # 진행 상황 업데이트
    progress_bar.progress(uploaded_count / total_files)
    progress_container.info(f"진행 상황: {uploaded_count}/{total_files} 완료")
    
    # 다음 단계로 버튼 - DB에서 확인하도록 수정
    if st.button("다음 단계: 접수 완료", key="next_to_complete"):
        # DB에서 직접 파일 및 사유 정보 확인
        conn = sqlite3.connect('audit_system.db')
        c = conn.cursor()
        
        # 파일명에 파일 유형 포함여부 확인
        incomplete_files = []
        for req_file in required_files:
            # 업로드 파일 확인
            c.execute("SELECT COUNT(*) FROM uploaded_files WHERE submission_id = ? AND file_name LIKE ?", 
                    (submission_id, f"%{req_file}%"))
            file_count = c.fetchone()[0]
            
            # 사유 제공 확인
            c.execute("SELECT COUNT(*) FROM missing_file_reasons WHERE submission_id = ? AND file_name = ?", 
                    (submission_id, req_file))
            reason_count = c.fetchone()[0]
            
            if file_count == 0 and reason_count == 0:
                incomplete_files.append(req_file)
        
        conn.close()
        current_missing_files = incomplete_files
        
        if incomplete_files:
            st.warning("다음 파일이 필요합니다:\n- " + "\n- ".join(incomplete_files))
        else:
            # 페이지 전환
            st.query_params["menu"] = "접수 완료"
            st.rerun()
      
# 접수 완료 페이지
elif menu == "접수 완료":
    st.title("✅ 일상감사 접수 완료")

    # ─── DB에서 접수 정보 불러오기 ───
    sub_id = st.session_state["submission_id"]
    conn = sqlite3.connect('audit_system.db')
    c = conn.cursor()
    c.execute("""
        SELECT department, manager, phone, contract_name, contract_date, contract_amount
        FROM submissions
        WHERE submission_id = ?
    """, (sub_id,))
    result = c.fetchone()
    if result:
        department, manager, phone, contract_name, contract_date, contract_amount = result
    else:
        st.error("접수 정보를 찾을 수 없습니다. 파일 업로드 페이지에서 접수 정보를 먼저 입력해주세요.")
        department, manager, phone, contract_name, contract_date, contract_amount = "", "", "", "", "", ""

    # 접수 내용 요약
    st.markdown("### 접수 내용 요약")

    # 업로드된 파일 목록
    uploaded_file_list = []
    c.execute(
        "SELECT file_name, file_path FROM uploaded_files WHERE submission_id = ?",
        (sub_id,)
    )
    uploaded_db_files = c.fetchall()

    if uploaded_db_files:
        st.markdown("#### 업로드된 파일")
        for file_name, file_path in uploaded_db_files:
            st.success(f"✅ {file_name}")
            uploaded_file_list.append(file_path)

    # 누락된 파일 및 사유
    c.execute(
        "SELECT file_name, reason FROM missing_file_reasons WHERE submission_id = ?",
        (sub_id,)
    )
    missing_db_files = c.fetchall()
    
    if missing_db_files:
        st.markdown("#### 누락된 파일 및 사유")
        for file_name, reason in missing_db_files:
            st.info(f"📝 {file_name}: {reason}")

    # DB에서 누락 파일 확인 - 파일 유형으로 검색
    incomplete_files = []
    for req_file in required_files:
        # 업로드 파일 확인
        c.execute("SELECT COUNT(*) FROM uploaded_files WHERE submission_id = ? AND file_name LIKE ?", 
                  (sub_id, f"%{req_file}%"))
        file_count = c.fetchone()[0]
        
        # 사유 제공 확인
        c.execute("SELECT COUNT(*) FROM missing_file_reasons WHERE submission_id = ? AND file_name = ?", 
                  (sub_id, req_file))
        reason_count = c.fetchone()[0]
        if file_count == 0 and reason_count == 0:
            incomplete_files.append(req_file)
    current_missing_files = incomplete_files

# 이메일 발송 섹션
    st.markdown("### 이메일 발송")
    recipient_email = st.text_input("수신자 이메일 주소", value=to_email)
    email_subject = st.text_input("이메일 제목", value=f"일상감사 접수: {submission_id}")
    additional_message = st.text_area("추가 메시지", value="")

    # ✅ 버튼도 여기 안에 있어야 함
    if st.button('접수 완료 및 이메일 발송'):
        if current_missing_files:
            st.warning(f"누락된 파일: {', '.join(current_missing_files)}. 업로드 또는 사유를 입력해 주세요.")
        else:
            # 업로드된 파일들을 ZIP으로 압축
            zip_file_path = None
            if uploaded_file_list:
                zip_folder = os.path.join(base_folder, "zips")
                if not os.path.exists(zip_folder):
                    os.makedirs(zip_folder)
                
                zip_file_path = os.path.join(zip_folder, f"일상감사_파일_{submission_id}.zip")
                
                with zipfile.ZipFile(zip_file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for file_path in uploaded_file_list:
                        if os.path.exists(file_path):
                            zipf.write(file_path, os.path.basename(file_path))
                
                # ZIP 파일 다운로드 버튼 제공
                with open(zip_file_path, "rb") as f:
                    zip_data = f.read()
                    st.download_button(
                        label="모든 파일 다운로드 (ZIP)",
                        data=zip_data,
                        file_name=f"일상감사_파일_{submission_id}.zip",
                        mime="application/zip"
                    )
            
            # 이메일 첨부 파일 목록 준비
            email_attachments = []
            
            # ZIP 파일이 있으면 첨부
            if zip_file_path and os.path.exists(zip_file_path):
                email_attachments.append(zip_file_path)
            else:
                # ZIP 파일이 없으면 개별 파일 첨부
                email_attachments.extend(uploaded_file_list)
            
            # 이메일 본문 작성
            body = f"일상감사 접수 ID: {submission_id}\n"
            body += f"접수일자: {upload_date}\n\n"
            
            if additional_message:
                body += f"추가 메시지:\n{additional_message}\n\n"
            
            # 업로드된 파일 목록 추가
            body += "업로드된 파일 목록:\n"
            for file_name, _ in uploaded_db_files:
                body += f"- {file_name}\n"
            
            # 누락된 파일 및 사유 추가
            if missing_db_files:
                body += "\n누락된 파일 및 사유:\n"
                for file_name, reason in missing_db_files:
                    body += f"- {file_name} (사유: {reason})\n"
            
            # 첨부 파일 안내 추가
            if zip_file_path:
                body += "\n* 업로드된 파일들이 ZIP 파일로 압축되어 첨부되어 있습니다.\n"
            # ✅ [여기] GPT 보고서 생성 및 첨부 추가
            report_path = generate_audit_report_with_gpt(
                submission_id=submission_id,
                department=st.session_state.get("department", ""),
                manager=st.session_state.get("manager", ""),
                phone=st.session_state.get("phone", ""),
                contract_name=st.session_state.get("contract_name", ""),
                contract_date=st.session_state.get("contract_date", ""),
                contract_amount=st.session_state.get("contract_amount_formatted", ""),
                uploaded_files=[f for f, _ in uploaded_db_files],
                missing_files_with_reasons=[(f, r) for f, r in missing_db_files]
            )

            if report_path and os.path.exists(report_path):
                email_attachments.append(report_path)
                body += "* GPT 기반 감사보고서 초안이 첨부되어 있습니다.\n"
            # 이메일 발송
            with st.spinner("이메일을 발송 중입니다..."):
                success, message = send_email(email_subject, body, recipient_email, email_attachments)
                
                if success:
                    # 데이터베이스에 접수 상태 업데이트
                    update_submission_status(submission_id, "접수완료", 1)
                    st.success("일상감사 접수가 완료되었으며, 이메일 알림이 발송되었습니다!")
                    
                    # 접수 완료 확인서 표시
                    st.markdown("### 접수 완료 확인서")
                    st.markdown(f"""
                    **접수 ID**: {submission_id}  
                    **접수일자**: {upload_date}  
                    **처리상태**: 접수완료  
                    **이메일 발송**: 완료 ({recipient_email})
                    """)
                    
                    # 다운로드 버튼 제공
                    receipt_text = f"""
                    일상감사 접수 확인서
                    
                    접수 ID: {submission_id}
                    접수일자: {upload_date}
                    처리상태: 접수완료
                    이메일 발송: 완료 ({recipient_email})
                    
                    업로드된 파일 목록:
                    """
                    for file_name, _ in uploaded_db_files:
                        receipt_text += f"- {file_name}\n"
                    
                    if missing_db_files:
                        receipt_text += "\n누락된 파일 및 사유:\n"
                        for file_name, reason in missing_db_files:
                            receipt_text += f"- {file_name} (사유: {reason})\n"
                    
                    st.download_button(
                        label="접수 확인서 다운로드",
                        data=receipt_text,
                        file_name=f"접수확인서_{submission_id}.txt",
                        mime="text/plain"
                    )
                    
                    # 이메일 발송 후 메모리 정리
                    for attachment in email_attachments:
                        if os.path.exists(attachment):
                            try:
                                # ZIP 파일은 남기고 나머지는 삭제 (선택적)
                                if not attachment.endswith('.zip'):
                                    os.remove(attachment)
                            except Exception as e:
                                logger.error(f"첨부파일 정리 오류: {str(e)}")
                    
                    # 캐시 데이터 초기화
                    st.cache_data.clear()
                    gc.collect()
                else:
                    st.error(f"이메일 발송 중 오류가 발생했습니다: {message}")


# 페이지 하단 정보
st.sidebar.markdown("---")
st.sidebar.info("""
2025 일상감사 접수 시스템
문의:  
    OKH. 감사팀
    📞 02-2009-6512/ 신승식
""")
